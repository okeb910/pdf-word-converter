import io
import tempfile
import unittest
import zipfile
from pathlib import Path

import pymupdf
from PIL import Image
from docx import Document
from docx.oxml.ns import qn

from pdf_word_converter import calculate_adaptive_dpi, pdf_to_word_via_images


class ImageDocxLayoutTests(unittest.TestCase):
    def test_image_docx_preserves_vector_details_as_uncompressed_rgb_png(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            folder = Path(temp_dir)
            source = folder / "anonymous-vector-form.pdf"
            output = folder / "anonymous-vector-form.docx"

            pdf = pymupdf.open()
            try:
                page = pdf.new_page(width=288, height=216)
                for x in (36, 108, 180):
                    page.draw_line((x, 36), (x, 120), color=(0, 0, 0), width=2)
                for y in (36, 78, 120):
                    page.draw_line((36, y), (180, y), color=(0, 0, 0), width=2)
                page.draw_line((208, 62), (220, 74), color=(0, 0, 0), width=4)
                page.draw_line((220, 74), (252, 42), color=(0, 0, 0), width=4)
                pdf.save(source)
            finally:
                pdf.close()

            pdf_to_word_via_images(
                source,
                output,
                lambda _message, _pct: None,
                dpi=72,
            )

            document = Document(output)
            settings = document.settings.element
            self.assertIsNotNone(
                settings.find(qn("w:doNotAutoCompressPictures"))
            )

            with zipfile.ZipFile(output) as archive:
                media_names = [
                    name
                    for name in archive.namelist()
                    if name.startswith("word/media/")
                ]
                self.assertEqual(len(media_names), 1)
                image_bytes = archive.read(media_names[0])

            with Image.open(io.BytesIO(image_bytes)) as image:
                self.assertEqual(image.format, "PNG")
                self.assertEqual(image.mode, "RGB")
                self.assertNotIn("A", image.getbands())

                grayscale = image.convert("L")
                table_region = grayscale.crop((32, 32, 184, 124))
                check_region = grayscale.crop((202, 36, 258, 80))
                self.assertGreater(
                    sum(pixel < 245 for pixel in table_region.getdata()),
                    500,
                )
                self.assertGreater(
                    sum(pixel < 245 for pixel in check_region.getdata()),
                    75,
                )

    def test_adaptive_dpi_obeys_dimension_and_total_pixel_limits(self):
        cases = (
            {
                "name": "long_edge_limit",
                "width_points": 7200,
                "height_points": 720,
                "expected_dpi": 60,
            },
            {
                "name": "total_pixel_limit",
                "width_points": 1440,
                "height_points": 1440,
                "expected_dpi": 273,
            },
        )
        for case in cases:
            with self.subTest(case["name"]):
                dpi = calculate_adaptive_dpi(
                    case["width_points"],
                    case["height_points"],
                    base_dpi=300,
                    max_dimension_px=6000,
                    max_pixels=30_000_000,
                )

                rendered_width = case["width_points"] / 72 * dpi
                rendered_height = case["height_points"] / 72 * dpi
                self.assertEqual(dpi, case["expected_dpi"])
                self.assertLessEqual(max(rendered_width, rendered_height), 6000)
                self.assertLessEqual(
                    rendered_width * rendered_height,
                    30_000_000,
                )

        with self.assertRaisesRegex(ValueError, "1 DPI"):
            calculate_adaptive_dpi(
                1_000_000,
                1_000_000,
                base_dpi=300,
                max_dimension_px=6000,
                max_pixels=30_000_000,
            )

    def test_mixed_page_images_are_placed_after_section_breaks(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            folder = Path(temp_dir)
            source = folder / "中文 mixed pages.pdf"
            output = folder / "中文 mixed pages.docx"

            pdf = pymupdf.open()
            pdf.new_page(width=595, height=842)
            pdf.new_page(width=842, height=595)
            pdf.save(source)
            pdf.close()

            pdf_to_word_via_images(source, output, lambda _message, _pct: None, dpi=72)

            document = Document(output)
            self.assertEqual(len(document.sections), 2)
            self.assertEqual(len(document.inline_shapes), 2)
            self.assertEqual(len(document.paragraphs), 3)

            section_break = document.paragraphs[1]._p
            second_image = document.paragraphs[2]._p
            self.assertTrue(section_break.xpath("./w:pPr/w:sectPr"))
            self.assertFalse(section_break.xpath(".//w:drawing"))
            self.assertTrue(second_image.xpath(".//w:drawing"))
            self.assertFalse(second_image.xpath("./w:pPr/w:sectPr"))

            self.assertLess(
                document.sections[0].page_width,
                document.sections[0].page_height,
            )
            self.assertGreater(
                document.sections[1].page_width,
                document.sections[1].page_height,
            )


if __name__ == "__main__":
    unittest.main()
