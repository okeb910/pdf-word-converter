import tempfile
import unittest
from dataclasses import FrozenInstanceError
from pathlib import Path
from unittest.mock import MagicMock, patch
from zipfile import ZipFile
from docx import Document
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

import pymupdf

from pdf_fidelity import (
    PdfFidelityAnalysisCancelled,
    PdfFidelityRisk,
    _deduplicate_table_candidates,
    _line_segments_cover_edge,
    _table_cell_border_edges_from_drawings,
    _document_has_tagged_table_structure,
    _select_precise_line_tables,
    _text_mapping_warning_count,
    analyze_pdf_fidelity_risk,
    inspect_editable_docx_tables,
)


class PdfFidelityTests(unittest.TestCase):
    @staticmethod
    def _save_blank_pdf(path: Path, page_count: int = 1) -> None:
        document = pymupdf.open()
        try:
            for _ in range(page_count):
                document.new_page(width=595, height=842)
            document.save(path)
        finally:
            document.close()

    @staticmethod
    def _draw_table(page) -> None:
        left, top, cell_width, cell_height = 72, 90, 110, 42
        rows, columns = 3, 3
        for row in range(rows + 1):
            y = top + row * cell_height
            page.draw_line(
                (left, y),
                (left + columns * cell_width, y),
                color=(0, 0, 0),
                width=1,
            )
        for column in range(columns + 1):
            x = left + column * cell_width
            page.draw_line(
                (x, top),
                (x, top + rows * cell_height),
                color=(0, 0, 0),
                width=1,
            )
        for row in range(rows):
            for column in range(columns):
                page.insert_text(
                    (left + column * cell_width + 8, top + row * cell_height + 25),
                    f"R{row + 1}C{column + 1}",
                    fontsize=10,
                )

    @staticmethod
    def _draw_borderless_table(page) -> None:
        for row in range(5):
            for column in range(3):
                page.insert_text(
                    (
                        72 + column * 150,
                        100 + row * 30,
                    ),
                    f"R{row + 1}C{column + 1}",
                    fontsize=10,
                )

    @staticmethod
    def _set_explicit_cell_borders(cell, **edge_values) -> None:
        cell_element = getattr(cell, "_tc", cell)
        properties = cell_element.get_or_add_tcPr()
        borders = properties.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            properties.append(borders)
        for name in ("top", "left", "bottom", "right"):
            existing = borders.find(qn(f"w:{name}"))
            if existing is not None:
                borders.remove(existing)
            value = edge_values.get(name)
            if value is None:
                continue
            edge = OxmlElement(f"w:{name}")
            edge.set(qn("w:val"), str(value))
            if str(value).casefold() not in {"nil", "none"}:
                edge.set(qn("w:sz"), "4")
                edge.set(qn("w:space"), "0")
                edge.set(qn("w:color"), "000000")
            borders.append(edge)

    def test_blank_pdf_is_simple_and_supports_chinese_space_path(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "中文 空格 空白.pdf"
            self._save_blank_pdf(source)

            result = analyze_pdf_fidelity_risk(source)

        self.assertIsInstance(result, PdfFidelityRisk)
        self.assertEqual(result.path, source)
        self.assertEqual((result.page_count, result.analyzed_pages), (1, 1))
        self.assertEqual(result.table_count, 0)
        self.assertEqual(result.table_text_character_count, 0)
        self.assertEqual(result.table_texts, ())
        self.assertEqual(result.table_cell_matrices, ())
        self.assertEqual(result.table_shapes, ())
        self.assertEqual(result.table_cell_counts, ())
        self.assertEqual(result.table_cell_border_edges, ())
        self.assertEqual(result.table_text_extraction_failure_count, 0)
        self.assertEqual(result.selectable_text_character_count, 0)
        self.assertEqual(result.selectable_text, "")
        self.assertEqual(result.unverifiable_text_character_count, 0)
        self.assertFalse(result.table_analysis_limited)
        self.assertFalse(result.table_analysis_uncertain)
        self.assertFalse(result.editable_table_candidate)
        self.assertEqual(result.widget_count, 0)
        self.assertEqual(result.vector_mark_count, 0)
        self.assertEqual(result.vector_path_count, 0)
        self.assertFalse(result.is_complex)
        self.assertEqual(result.reasons, ())
        with self.assertRaises(FrozenInstanceError):
            result.is_complex = True

    def test_unverifiable_text_mapping_characters_are_counted(self):
        self.assertEqual(_text_mapping_warning_count("正常文字ABC"), 0)
        self.assertEqual(_text_mapping_warning_count("\ufffd\ue000"), 2)

    def test_table_and_many_vector_paths_are_reported(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "anonymous-table.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                self._draw_table(page)
                for index in range(30):
                    y = 300 + index * 4
                    page.draw_line((72, y), (200, y), color=(0, 0, 0), width=0.5)
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(source, vector_threshold=20)

        self.assertGreaterEqual(result.table_count, 1)
        self.assertGreater(result.table_text_character_count, 0)
        self.assertEqual(len(result.table_texts), result.table_count)
        self.assertIn((3, 3), result.table_shapes)
        self.assertIn(9, result.table_cell_counts)
        self.assertIn("R1C1", "".join(result.table_texts))
        expected_matrix = tuple(
            tuple(
                f"R{row + 1}C{column + 1}"
                for column in range(3)
            )
            for row in range(3)
        )
        self.assertIn(expected_matrix, result.table_cell_matrices)
        table_index = result.table_shapes.index((3, 3))
        self.assertEqual(len(result.table_cell_border_edges[table_index]), 9)
        self.assertTrue(
            all(
                edges == (True, True, True, True)
                for edges in result.table_cell_border_edges[table_index]
            )
        )
        self.assertTrue(result.editable_table_candidate)
        self.assertEqual(result.table_text_extraction_failure_count, 0)
        self.assertGreater(result.selectable_text_character_count, 0)
        self.assertTrue(result.has_selectable_text)
        self.assertGreater(result.vector_path_count, 20)
        self.assertTrue(result.is_complex)
        self.assertTrue(any("表格" in reason for reason in result.reasons))
        self.assertTrue(any("矢量路径" in reason for reason in result.reasons))

    def test_borderless_aligned_text_table_is_detected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "borderless-table.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                self._draw_borderless_table(page)
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(
                source,
                max_pages=None,
            )

        self.assertEqual(result.table_count, 1)
        self.assertEqual(result.vector_path_count, 0)
        self.assertEqual(result.table_shapes, ((5, 3),))
        self.assertEqual(result.table_cell_counts, (15,))
        self.assertEqual(
            result.table_cell_border_edges,
            (((False, False, False, False),) * 15,),
        )
        self.assertGreater(result.table_text_character_count, 0)
        self.assertEqual(len(result.table_texts), result.table_count)
        self.assertIn("R1C1", "".join(result.table_texts))
        self.assertTrue(result.editable_table_candidate)
        self.assertTrue(result.has_selectable_text)

    def test_rotated_pdf_border_evidence_allows_three_point_geometry_drift(self):
        table = MagicMock()
        row = MagicMock()
        row.cells = ((101.81, 514.0, 167.684, 529.0),)
        table.rows = (row,)
        page = MagicMock()
        page.rotation = 90
        page.derotation_matrix = pymupdf.Matrix(0, -1, 1, 0, 0, 842)
        raw_rect = pymupdf.Rect(514.0, 671.55, 529.0, 740.19)
        drawing = {
            "type": "s",
            "color": (0.0, 0.0, 0.0),
            "stroke_opacity": 1.0,
            "width": 0.5,
            "items": (
                ("l", raw_rect.tl, raw_rect.tr),
                ("l", raw_rect.tr, raw_rect.br),
                ("l", raw_rect.br, raw_rect.bl),
                ("l", raw_rect.bl, raw_rect.tl),
            ),
        }

        edges = _table_cell_border_edges_from_drawings(
            page,
            table,
            ((0, 0, 1, 1),),
            (drawing,),
        )

        self.assertEqual(edges, ((True, True, True, True),))

    def test_pdf_border_evidence_supports_short_edges_and_rejects_gaps(self):
        table = MagicMock()
        row = MagicMock()
        row.cells = ((10.0, 10.0, 18.0, 18.0),)
        table.rows = (row,)
        page = MagicMock()
        page.rotation = 0
        drawing = {
            "type": "s",
            "color": (0.0, 0.0, 0.0),
            "stroke_opacity": 1.0,
            "width": 0.5,
            "items": (
                ("l", pymupdf.Point(10, 10), pymupdf.Point(13, 10)),
                ("l", pymupdf.Point(15, 10), pymupdf.Point(18, 10)),
                ("l", pymupdf.Point(10, 10), pymupdf.Point(10, 18)),
                ("l", pymupdf.Point(10, 18), pymupdf.Point(18, 18)),
                ("l", pymupdf.Point(18, 10), pymupdf.Point(18, 18)),
            ),
        }

        edges = _table_cell_border_edges_from_drawings(
            page,
            table,
            ((0, 0, 1, 1),),
            (drawing,),
        )

        self.assertEqual(edges, ((False, True, True, True),))

    def test_pdf_border_evidence_rejects_short_middle_fragment(self):
        table = MagicMock()
        row = MagicMock()
        row.cells = ((10.0, 10.0, 18.0, 18.0),)
        table.rows = (row,)
        page = MagicMock()
        page.rotation = 0
        drawing = {
            "type": "s",
            "color": (0.0, 0.0, 0.0),
            "stroke_opacity": 1.0,
            "width": 0.5,
            "items": (
                ("l", pymupdf.Point(13, 10), pymupdf.Point(15, 10)),
            ),
        }

        edges = _table_cell_border_edges_from_drawings(
            page,
            table,
            ((0, 0, 1, 1),),
            (drawing,),
        )

        self.assertEqual(edges, ((False, False, False, False),))

    def test_pdf_border_evidence_accepts_thin_fill_only_rectangles(self):
        table = MagicMock()
        row = MagicMock()
        row.cells = ((10.0, 10.0, 18.0, 18.0),)
        table.rows = (row,)
        page = MagicMock()
        page.rotation = 0
        drawing = {
            "type": "f",
            "fill": (0.0, 0.0, 0.0),
            "fill_opacity": 1.0,
            "items": (
                ("re", pymupdf.Rect(10.0, 9.5, 18.0, 10.5)),
                ("re", pymupdf.Rect(9.5, 10.0, 10.5, 18.0)),
                ("re", pymupdf.Rect(10.0, 17.5, 18.0, 18.5)),
                ("re", pymupdf.Rect(17.5, 10.0, 18.5, 18.0)),
            ),
        }

        edges = _table_cell_border_edges_from_drawings(
            page,
            table,
            ((0, 0, 1, 1),),
            (drawing,),
        )

        self.assertEqual(edges, ((True, True, True, True),))

    def test_pdf_border_evidence_accepts_real_pymupdf_fill_only_rule(self):
        document = pymupdf.open()
        try:
            page = document.new_page(width=100, height=100)
            page.draw_rect(
                pymupdf.Rect(10.0, 9.5, 18.0, 10.5),
                color=None,
                fill=(0.0, 0.0, 0.0),
            )
            table = MagicMock()
            row = MagicMock()
            row.cells = ((10.0, 10.0, 18.0, 18.0),)
            table.rows = (row,)

            edges = _table_cell_border_edges_from_drawings(
                page,
                table,
                ((0, 0, 1, 1),),
                page.get_drawings(),
            )
        finally:
            document.close()

        self.assertEqual(edges, ((True, False, False, False),))

    def test_pdf_border_evidence_requires_seventy_percent_coverage(self):
        edge = ((0.0, 0.0), (10.0, 0.0))
        self.assertTrue(
            _line_segments_cover_edge(
                *edge,
                (((1.5, 0.0), (8.5, 0.0)),),
                10.0,
            )
        )
        self.assertFalse(
            _line_segments_cover_edge(
                *edge,
                (((1.51, 0.0), (8.49, 0.0)),),
                10.0,
            )
        )

    def test_pdf_border_evidence_ignores_block_fill_and_transparent_paths(self):
        table = MagicMock()
        row = MagicMock()
        row.cells = ((10.0, 10.0, 18.0, 18.0),)
        table.rows = (row,)
        page = MagicMock()
        page.rotation = 0
        item = ("re", pymupdf.Rect(10.0, 10.0, 18.0, 18.0))
        drawings = (
            {
                "type": "f",
                "fill": (0.0, 0.0, 0.0),
                "fill_opacity": 1.0,
                "items": (item,),
            },
            {
                "type": "s",
                "color": (0.0, 0.0, 0.0),
                "stroke_opacity": 0.0,
                "width": 0.5,
                "items": (item,),
            },
        )

        edges = _table_cell_border_edges_from_drawings(
            page,
            table,
            ((0, 0, 1, 1),),
            drawings,
        )

        self.assertEqual(edges, ((False, False, False, False),))

    def test_two_row_borderless_aligned_text_table_is_detected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "two-row-borderless-table.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                for row in range(2):
                    for column in range(3):
                        page.insert_text(
                            (
                                72 + column * 150,
                                100 + row * 30,
                            ),
                            f"R{row + 1}C{column + 1}",
                            fontsize=10,
                        )
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(
                source,
                max_pages=None,
            )

        self.assertEqual(result.table_count, 1)
        self.assertGreater(result.table_text_character_count, 0)
        self.assertEqual(len(result.table_texts), result.table_count)
        self.assertIn("R1C1", "".join(result.table_texts))
        self.assertTrue(result.editable_table_candidate)
        self.assertTrue(result.has_selectable_text)

    def test_independent_two_column_row_flags_empty_primary_table_result(self):
        page = MagicMock()
        page.rect = pymupdf.Rect(0, 0, 595, 842)
        words = (
            (72, 390, 100, 404, "Owner:", 0, 0, 0),
            (112, 390, 160, 404, "Alice", 0, 0, 1),
        )
        text_dictionary = {
            "blocks": [
                {
                    "lines": [
                        {
                            "spans": [
                                {
                                    "text": "Owner: Alice",
                                    "font": "Arial",
                                }
                            ]
                        }
                    ]
                }
            ]
        }
        page.get_text.side_effect = (
            lambda mode, *args, **kwargs: (
                text_dictionary if mode == "dict" else words if mode == "words" else ""
            )
        )
        page.find_tables.return_value.tables = ()
        page.widgets.return_value = None
        page.get_drawings.return_value = []
        document = MagicMock()
        document.__len__.return_value = 1
        document.__getitem__.return_value = page

        with patch("pdf_fidelity.pymupdf.open", return_value=document):
            result = analyze_pdf_fidelity_risk("unconfirmed-two-column-row.pdf")

        self.assertEqual(result.table_count, 0)
        self.assertTrue(result.table_layout_suspected)
        self.assertTrue(result.has_selectable_text)
        self.assertTrue(any("表格结构无法确认" in reason for reason in result.reasons))
        document.close.assert_called_once_with()

    def test_independent_text_layout_flags_empty_primary_table_result(self):
        page = MagicMock()
        page.rect = pymupdf.Rect(0, 0, 595, 842)
        words = tuple(
            (
                72 + column * 150,
                90 + row * 30,
                110 + column * 150,
                102 + row * 30,
                f"R{row + 1}C{column + 1}",
                0,
                row,
                column,
            )
            for row in range(2)
            for column in range(3)
        )
        text_dictionary = {
            "blocks": [
                {
                    "lines": [
                        {
                            "spans": [
                                {
                                    "text": "R1C1 R1C2 R1C3 R2C1 R2C2 R2C3",
                                    "font": "Arial",
                                }
                            ]
                        }
                    ]
                }
            ]
        }
        page.get_text.side_effect = (
            lambda mode, *args, **kwargs: (
                text_dictionary if mode == "dict" else words if mode == "words" else ""
            )
        )
        page.find_tables.return_value.tables = ()
        page.widgets.return_value = None
        page.get_drawings.return_value = []
        document = MagicMock()
        document.__len__.return_value = 1
        document.__getitem__.return_value = page

        with patch("pdf_fidelity.pymupdf.open", return_value=document):
            result = analyze_pdf_fidelity_risk("unconfirmed-text-table.pdf")

        self.assertEqual(result.table_count, 0)
        self.assertFalse(result.editable_table_candidate)
        self.assertTrue(result.table_layout_suspected)
        self.assertTrue(result.has_selectable_text)
        self.assertTrue(result.is_complex)
        self.assertTrue(any("表格结构无法确认" in reason for reason in result.reasons))
        document.close.assert_called_once_with()

    def test_independent_grid_layout_flags_empty_primary_table_result(self):
        page = MagicMock()
        page.rect = pymupdf.Rect(0, 0, 595, 842)
        text_dictionary = {
            "blocks": [
                {
                    "lines": [
                        {
                            "spans": [
                                {
                                    "text": "Grid values",
                                    "font": "Arial",
                                }
                            ]
                        }
                    ]
                }
            ]
        }
        page.get_text.side_effect = (
            lambda mode, *args, **kwargs: (
                text_dictionary if mode == "dict" else () if mode == "words" else ""
            )
        )
        drawings = []
        for y_position in (90, 130, 170):
            drawings.append(
                {
                    "items": [
                        (
                            "l",
                            pymupdf.Point(72, y_position),
                            pymupdf.Point(292, y_position),
                        )
                    ]
                }
            )
        for x_position in (72, 182, 292):
            drawings.append(
                {
                    "items": [
                        (
                            "l",
                            pymupdf.Point(x_position, 90),
                            pymupdf.Point(x_position, 170),
                        )
                    ]
                }
            )
        page.find_tables.return_value.tables = ()
        page.widgets.return_value = None
        page.get_drawings.return_value = drawings
        document = MagicMock()
        document.__len__.return_value = 1
        document.__getitem__.return_value = page

        with patch("pdf_fidelity.pymupdf.open", return_value=document):
            result = analyze_pdf_fidelity_risk("unconfirmed-grid-table.pdf")

        self.assertEqual(result.table_count, 0)
        self.assertTrue(result.table_layout_suspected)
        self.assertGreater(result.selectable_text_character_count, 0)
        self.assertTrue(any("表格结构无法确认" in reason for reason in result.reasons))
        document.close.assert_called_once_with()

    def test_blank_line_and_rectangle_grids_fail_closed(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            folder = Path(temp_dir)
            for drawing_kind in ("line", "rectangle"):
                with self.subTest(drawing_kind=drawing_kind):
                    source = folder / f"blank-{drawing_kind}-grid.pdf"
                    document = pymupdf.open()
                    try:
                        page = document.new_page(width=595, height=842)
                        if drawing_kind == "line":
                            for y_position in (180, 230, 280):
                                page.draw_line(
                                    (72, y_position),
                                    (292, y_position),
                                    color=(0, 0, 0),
                                    width=1,
                                )
                            for x_position in (72, 182, 292):
                                page.draw_line(
                                    (x_position, 180),
                                    (x_position, 280),
                                    color=(0, 0, 0),
                                    width=1,
                                )
                        else:
                            for row in range(2):
                                for column in range(2):
                                    page.draw_rect(
                                        pymupdf.Rect(
                                            72 + column * 110,
                                            180 + row * 50,
                                            182 + column * 110,
                                            230 + row * 50,
                                        ),
                                        color=(0, 0, 0),
                                        width=1,
                                    )
                        document.save(source)
                    finally:
                        document.close()

                    result = analyze_pdf_fidelity_risk(
                        source,
                        max_pages=None,
                    )

                self.assertEqual(
                    result.selectable_text_character_count,
                    0,
                )
                self.assertFalse(result.editable_table_candidate)
                self.assertTrue(result.table_layout_suspected)
                self.assertTrue(result.is_complex)
                self.assertTrue(
                    any(
                        "表格结构无法确认" in reason
                        for reason in result.reasons
                    )
                )

    def test_blank_quadrilateral_grid_fails_closed(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "blank-quadrilateral-grid.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                for row in range(2):
                    for column in range(2):
                        cell = pymupdf.Rect(
                            72 + column * 110,
                            180 + row * 50,
                            182 + column * 110,
                            230 + row * 50,
                        )
                        page.draw_quad(
                            cell.quad,
                            color=(0, 0, 0),
                            width=1,
                        )
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(source, max_pages=None)

        self.assertEqual(result.selectable_text_character_count, 0)
        self.assertFalse(result.editable_table_candidate)
        self.assertTrue(result.table_layout_suspected)
        self.assertTrue(result.is_complex)
        self.assertTrue(
            any(
                "表格结构无法确认" in reason
                for reason in result.reasons
            )
        )
    def test_cubic_bezier_grid_is_never_treated_as_plain_text(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "cubic-grid.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                shape = page.new_shape()
                for y_position in (180, 230):
                    shape.draw_bezier(
                        (72, y_position),
                        (145, y_position),
                        (219, y_position),
                        (292, y_position),
                    )
                for x_position in (72, 182, 292):
                    shape.draw_bezier(
                        (x_position, 180),
                        (x_position, 197),
                        (x_position, 213),
                        (x_position, 230),
                    )
                shape.finish(color=(0, 0, 0), width=1)
                shape.commit()
                page.insert_text((82, 210), "Left", fontsize=10)
                page.insert_text((192, 210), "Right", fontsize=10)
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(source, max_pages=None)

        self.assertGreater(result.selectable_text_character_count, 0)
        self.assertTrue(result.table_count or result.table_layout_suspected)
        self.assertTrue(
            result.editable_table_candidate or result.table_layout_suspected
        )

    def test_tagged_table_without_geometric_detection_fails_closed(self):
        page = MagicMock()
        page.rect = pymupdf.Rect(0, 0, 595, 842)
        page.get_text.side_effect = (
            lambda mode, *args, **kwargs: (
                {
                    "blocks": [
                        {
                            "lines": [
                                {
                                    "spans": [
                                        {"text": "Tagged cells", "font": "Arial"}
                                    ]
                                }
                            ]
                        }
                    ]
                }
                if mode == "dict"
                else ()
                if mode == "words"
                else ""
            )
        )
        page.find_tables.return_value.tables = ()
        page.widgets.return_value = None
        page.get_drawings.return_value = []

        document = MagicMock()
        document.__len__.return_value = 1
        document.__getitem__.return_value = page
        document.pdf_catalog.return_value = 1

        def xref_get_key(xref, key):
            values = {
                (1, "StructTreeRoot"): ("xref", "2 0 R"),
                (2, "RoleMap"): (
                    "dict",
                    "<< /Custom#54able /Table /CustomRow /TR /CustomCell /TD >>",
                ),
                (2, "S"): ("name", "/Custom#54able"),
                (2, "K"): ("null", "null"),
            }
            return values.get((xref, key), ("null", "null"))

        document.xref_get_key.side_effect = xref_get_key

        with patch("pdf_fidelity.pymupdf.open", return_value=document):
            result = analyze_pdf_fidelity_risk("tagged-table.pdf")

        self.assertTrue(result.tagged_table_structure_present)
        self.assertTrue(result.table_layout_suspected)
        self.assertEqual(result.table_count, 0)
        self.assertTrue(any("标签结构" in reason for reason in result.reasons))
        document.close.assert_called_once_with()

    def test_tagged_table_follows_indirect_k_array(self):
        document = MagicMock()
        document.pdf_catalog.return_value = 1

        def xref_get_key(xref, key):
            values = {
                (1, "StructTreeRoot"): ("xref", "2 0 R"),
                (2, "RoleMap"): ("null", "null"),
                (2, "S"): ("null", "null"),
                (2, "K"): ("xref", "3 0 R"),
                (4, "S"): ("name", "/Table"),
                (4, "K"): ("null", "null"),
            }
            return values.get((xref, key), ("null", "null"))

        object_values = {
            2: "<< /Type /StructTreeRoot /K 3 0 R >>",
            3: "[ 4 0 R ]",
            4: "<< /Type /StructElem /S /Table /P 2 0 R >>",
        }
        document.xref_get_key.side_effect = xref_get_key
        document.xref_object.side_effect = (
            lambda xref, compressed=False: object_values[xref]
        )

        self.assertTrue(_document_has_tagged_table_structure(document))

    def test_tagged_table_detects_inline_escaped_cell_role(self):
        document = MagicMock()
        document.pdf_catalog.return_value = 1

        def xref_get_key(xref, key):
            values = {
                (1, "StructTreeRoot"): ("xref", "2 0 R"),
                (2, "RoleMap"): ("null", "null"),
                (2, "S"): ("null", "null"),
                (2, "K"): (
                    "array",
                    "[ << /Type /StructElem /S /T#44 /K 0 >> ]",
                ),
            }
            return values.get((xref, key), ("null", "null"))

        document.xref_get_key.side_effect = xref_get_key
        document.xref_object.return_value = (
            "<< /Type /StructTreeRoot "
            "/K [ << /Type /StructElem /S /T#44 /K 0 >> ] >>"
        )

        self.assertTrue(_document_has_tagged_table_structure(document))

    def test_lined_table_keeps_a_real_empty_row_in_its_shape(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "lined-empty-row.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                left, top, cell_width, cell_height = 72, 90, 110, 42
                for row in range(4):
                    y = top + row * cell_height
                    page.draw_line((left, y), (left + 330, y), color=(0, 0, 0))
                for column in range(4):
                    x = left + column * cell_width
                    page.draw_line((x, top), (x, top + 126), color=(0, 0, 0))
                for row in (0, 2):
                    for column in range(3):
                        page.insert_text(
                            (left + column * cell_width + 8, top + row * cell_height + 25),
                            f"R{row + 1}C{column + 1}",
                            fontsize=10,
                        )
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(source, max_pages=None)

        self.assertEqual(result.table_shapes, ((3, 3),))
        self.assertEqual(result.table_cell_counts, (9,))
    def test_lined_and_borderless_tables_on_same_page_are_both_detected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "mixed-table-styles.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                self._draw_table(page)
                for row in range(3):
                    for column in range(3):
                        page.insert_text(
                            (72 + column * 150, 430 + row * 32),
                            f"B{row + 1}C{column + 1}",
                            fontsize=10,
                        )
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(source, max_pages=None)

        self.assertEqual(result.table_count, 2)
        self.assertIn((3, 3), result.table_shapes)
        self.assertEqual(result.table_shapes.count((3, 3)), 2)
        self.assertEqual(result.table_cell_counts, (9, 9))
        self.assertIn("R1C1", "".join(result.table_texts))
        self.assertIn("B1C1", "".join(result.table_texts))

    def test_mixed_table_candidates_are_sorted_in_page_reading_order(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "mixed-reading-order.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                for row in range(3):
                    for column in range(3):
                        page.insert_text(
                            (72 + column * 150, 100 + row * 30),
                            f"TOP{row + 1}{column + 1}",
                            fontsize=10,
                        )

                left, top, cell_width, cell_height = 72, 430, 110, 42
                for row in range(4):
                    y = top + row * cell_height
                    page.draw_line((left, y), (left + 330, y), color=(0, 0, 0))
                for column in range(4):
                    x = left + column * cell_width
                    page.draw_line((x, top), (x, top + 126), color=(0, 0, 0))
                for row in range(3):
                    for column in range(3):
                        page.insert_text(
                            (left + column * cell_width + 8, top + row * cell_height + 25),
                            f"BOTTOM{row + 1}{column + 1}",
                            fontsize=10,
                        )
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(source, max_pages=None)

        self.assertEqual(result.table_count, 2)
        self.assertIn("TOP", result.table_texts[0])
        self.assertIn("BOTTOM", result.table_texts[1])
        self.assertEqual(result.table_shapes, ((3, 3), (3, 3)))
    def test_middle_single_row_three_column_table_is_detected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "single-row-table.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                for column, value in enumerate(("Item-001", "Item-002", "Item-003")):
                    page.insert_text((72 + column * 180, 400), value, fontsize=10)
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(source, max_pages=None)

        self.assertEqual(result.table_count, 1)
        self.assertEqual(result.table_shapes, ((1, 3),))

    def test_single_row_headers_footers_ignored_but_middle_labels_are_suspected(self):
        cases = (
            (80, ("Header A", "Header B", "Header C")),
            (400, ("Owner:", "Reviewer:", "Date:")),
            (790, ("College", "Reviewer", "Print date")),
        )
        for y_position, values in cases:
            with self.subTest(y_position=y_position, values=values):
                with tempfile.TemporaryDirectory() as temp_dir:
                    source = Path(temp_dir) / "single-row-labels.pdf"
                    document = pymupdf.open()
                    try:
                        page = document.new_page(width=595, height=842)
                        for column, value in enumerate(values):
                            page.insert_text(
                                (72 + column * 180, y_position),
                                value,
                                fontsize=10,
                            )
                        document.save(source)
                    finally:
                        document.close()

                    result = analyze_pdf_fidelity_risk(source, max_pages=None)

                expected_table_count = 0 if y_position == 400 else 1
                self.assertEqual(result.table_count, expected_table_count)
                self.assertEqual(
                    result.editable_table_candidate,
                    bool(expected_table_count),
                )
                self.assertTrue(result.has_selectable_text)
                self.assertTrue(result.table_layout_suspected)

    def test_narrow_single_row_two_column_pdf_is_suspected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "narrow-single-row-table.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                page.insert_text((72, 400), "Item", fontsize=10)
                page.insert_text((105, 400), "Amount", fontsize=10)
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(source, max_pages=None)

        self.assertEqual(result.table_count, 0)
        self.assertTrue(result.table_layout_suspected)
        self.assertTrue(result.has_selectable_text)

    def test_compact_borderless_two_by_two_table_fails_closed(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "compact-borderless-two-by-two.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                left_x = 72
                right_x = (
                    left_x
                    + pymupdf.get_text_length("A1", fontsize=10)
                    + 4
                )
                for row, y_position in enumerate((380, 410), start=1):
                    page.insert_text(
                        (left_x, y_position),
                        f"A{row}",
                        fontsize=10,
                    )
                    page.insert_text(
                        (right_x, y_position),
                        f"B{row}",
                        fontsize=10,
                    )
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(source, max_pages=None)

        self.assertTrue(result.has_selectable_text)
        self.assertTrue(
            result.table_count or result.table_layout_suspected
        )
        self.assertTrue(
            result.editable_table_candidate
            or result.table_layout_suspected
        )

    def test_compact_borderless_two_by_nine_table_fails_closed(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "compact-borderless-two-by-nine.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                for row, y_position in enumerate((380, 410)):
                    for column in range(9):
                        page.insert_text(
                            (72 + column * 13, y_position),
                            chr(ord("A") + row + column),
                            fontsize=10,
                        )
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(source, max_pages=None)

        self.assertTrue(result.has_selectable_text)
        self.assertTrue(
            result.table_count or result.table_layout_suspected
        )

    def test_compact_table_with_merged_heading_fails_closed(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "compact-merged-heading.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                page.insert_text((72, 360), "Header", fontsize=10)
                page.insert_text((72, 390), "A", fontsize=10)
                page.insert_text((85, 390), "B", fontsize=10)
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(source, max_pages=None)

        self.assertTrue(result.has_selectable_text)
        self.assertTrue(
            result.table_count or result.table_layout_suspected
        )
    def test_normal_paragraph_is_not_misclassified_as_borderless_table(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "normal-paragraph.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                page.insert_textbox(
                    pymupdf.Rect(72, 72, 520, 500),
                    (
                        "This is a normal paragraph with several words and no "
                        "table structure. "
                    )
                    * 12,
                    fontsize=11,
                )
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(source, max_pages=None)

        self.assertEqual(result.table_count, 0)
        self.assertFalse(result.editable_table_candidate)
        self.assertTrue(result.has_selectable_text)
        self.assertFalse(result.is_complex)

    def test_small_vector_checkmark_is_reported(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "vector-checkmark-table.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                self._draw_table(page)
                page.draw_line(
                    (82, 108),
                    (87, 114),
                    color=(0, 0, 0),
                    width=1.2,
                )
                page.draw_line(
                    (87, 114),
                    (98, 98),
                    color=(0, 0, 0),
                    width=1.2,
                )
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(
                source,
                max_pages=None,
            )

        self.assertGreaterEqual(result.table_count, 1)
        self.assertGreaterEqual(result.vector_mark_count, 1)
        self.assertTrue(
            any(
                "矢量复选或勾选" in reason
                for reason in result.reasons
            )
        )
    def test_widgets_and_symbol_font_runs_are_reported(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "anonymous-form.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                widget = pymupdf.Widget()
                widget.field_name = "anonymous_choice"
                widget.field_type = pymupdf.PDF_WIDGET_TYPE_CHECKBOX
                widget.rect = pymupdf.Rect(72, 72, 90, 90)
                page.add_widget(widget)
                page.insert_text((110, 86), "a", fontname="symb", fontsize=12)
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(source)

        self.assertEqual(result.widget_count, 1)
        self.assertGreaterEqual(result.symbol_font_run_count, 1)
        self.assertTrue(result.is_complex)
        self.assertTrue(any("表单控件" in reason for reason in result.reasons))
        self.assertTrue(any("Symbol" in reason for reason in result.reasons))

    def test_explicit_checkbox_symbols_are_reported(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "anonymous-checkbox-symbols.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                page.insert_htmlbox(
                    pymupdf.Rect(72, 72, 500, 140),
                    (
                        "<p>Options: &#x2610; &#x2611; &#x2612; &#x2713; "
                        "&#x2714; &#x25A1; &#x221A;</p>"
                    ),
                )
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(source)

        self.assertEqual(result.checkbox_symbol_count, 7)
        self.assertTrue(result.is_complex)
        self.assertTrue(any("复选或勾选符号" in reason for reason in result.reasons))

    def test_max_pages_limits_analysis(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "anonymous-multipage.pdf"
            document = pymupdf.open()
            try:
                for page_number in range(4):
                    page = document.new_page(width=595, height=842)
                    if page_number == 3:
                        for index in range(12):
                            y = 100 + index * 8
                            page.draw_line((72, y), (300, y), color=(0, 0, 0))
                document.save(source)
            finally:
                document.close()

            limited = analyze_pdf_fidelity_risk(
                source, max_pages=3, vector_threshold=5
            )
            complete = analyze_pdf_fidelity_risk(
                source, max_pages=None, vector_threshold=5
            )

        self.assertEqual((limited.page_count, limited.analyzed_pages), (4, 3))
        self.assertEqual(limited.vector_path_count, 0)
        self.assertFalse(limited.is_complex)
        self.assertEqual((complete.page_count, complete.analyzed_pages), (4, 4))
        self.assertGreater(complete.vector_path_count, 5)
        self.assertTrue(complete.is_complex)

    def test_long_document_scans_all_text_but_uses_fast_table_analysis(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "anonymous-long-report.pdf"
            document = pymupdf.open()
            try:
                for page_number in range(21):
                    page = document.new_page(width=595, height=842)
                    page.insert_text((72, 72), f"Page {page_number + 1}")
                    if page_number == 20:
                        self._draw_table(page)
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(source, max_pages=None)

        self.assertEqual((result.page_count, result.analyzed_pages), (21, 21))
        self.assertTrue(result.table_analysis_limited)
        self.assertGreater(result.selectable_text_character_count, 0)
        self.assertIn("Page 21", result.selectable_text)
        self.assertEqual(result.table_count, 0)
        self.assertTrue(result.table_layout_suspected)
        self.assertGreater(result.unconfirmed_table_region_count, 0)
        self.assertTrue(
            any("快速文字检查" in reason for reason in result.reasons)
        )

    def test_detailed_table_page_limit_must_be_positive(self):
        with self.assertRaisesRegex(ValueError, "detailed_table_page_limit"):
            analyze_pdf_fidelity_risk(
                "unused.pdf",
                detailed_table_page_limit=0,
            )

    def test_analysis_error_is_clear_and_document_is_closed(self):
        page = MagicMock()
        page.find_tables.side_effect = ValueError("broken table data")
        document = MagicMock()
        document.__len__.return_value = 1
        document.__getitem__.return_value = page

        with patch("pdf_fidelity.pymupdf.open", return_value=document), self.assertRaisesRegex(
            RuntimeError,
            "无法分析 PDF.*broken table data",
        ):
            analyze_pdf_fidelity_risk("broken input.pdf")

        document.close.assert_called_once_with()

    def test_table_extraction_failure_uses_bbox_text_and_stays_protected(self):
        table = MagicMock()
        table.bbox = (10, 20, 300, 180)
        table.extract.side_effect = RuntimeError("cell extraction failed")
        page = MagicMock()
        page.find_tables.return_value.tables = (table,)
        page.get_textbox.return_value = "Header Value 100"
        page.get_text.return_value = {
            "blocks": [
                {
                    "lines": [
                        {
                            "spans": [
                                {
                                    "text": "Header Value 100",
                                    "font": "Arial",
                                }
                            ]
                        }
                    ]
                }
            ]
        }
        page.widgets.return_value = None
        page.get_drawings.return_value = []
        document = MagicMock()
        document.__len__.return_value = 1
        document.__getitem__.return_value = page

        with patch("pdf_fidelity.pymupdf.open", return_value=document):
            result = analyze_pdf_fidelity_risk("uncertain-table.pdf")

        self.assertEqual(result.table_count, 1)
        self.assertEqual(result.table_text_extraction_failure_count, 1)
        self.assertGreater(result.table_text_character_count, 0)
        self.assertGreater(result.selectable_text_character_count, 0)
        self.assertTrue(result.editable_table_candidate)
        self.assertTrue(result.table_analysis_uncertain)
        self.assertTrue(any("提取不完整" in reason for reason in result.reasons))
        document.close.assert_called_once_with()

    def test_empty_table_extraction_without_bbox_text_is_uncertain(self):
        table = MagicMock()
        table.bbox = None
        table.extract.return_value = []
        page = MagicMock()
        page.find_tables.return_value.tables = (table,)
        page.get_text.return_value = {
            "blocks": [{"lines": [{"spans": [{"text": "Table Value", "font": "Arial"}]}]}]
        }
        page.widgets.return_value = None
        page.get_drawings.return_value = []
        document = MagicMock()
        document.__len__.return_value = 1
        document.__getitem__.return_value = page

        with patch("pdf_fidelity.pymupdf.open", return_value=document):
            result = analyze_pdf_fidelity_risk("empty-table-extraction.pdf")

        self.assertEqual(result.table_count, 1)
        self.assertEqual(result.table_text_extraction_failure_count, 1)
        self.assertEqual(result.table_text_character_count, 0)
        self.assertEqual(result.table_texts, ("",))
        self.assertGreater(result.selectable_text_character_count, 0)
        self.assertTrue(result.editable_table_candidate)
        self.assertTrue(result.table_analysis_uncertain)
        self.assertTrue(result.reasons)
        document.close.assert_called_once_with()

    def test_partial_table_extraction_uses_bbox_text_baseline(self):
        table = MagicMock()
        table.bbox = (10, 20, 300, 180)
        table.row_count = 2
        table.col_count = 2
        table.extract.return_value = [["Alpha", "Beta"], [None, None]]
        page = MagicMock()
        page.find_tables.return_value.tables = (table,)
        page.get_textbox.return_value = "Alpha Beta Gamma"

        def get_text(kind):
            if kind == "dict":
                return {
                    "blocks": [
                        {
                            "lines": [
                                {
                                    "spans": [
                                        {"text": "Alpha Beta Gamma", "font": "Arial"}
                                    ]
                                }
                            ]
                        }
                    ]
                }
            if kind == "words":
                return []
            raise AssertionError(kind)

        page.get_text.side_effect = get_text
        page.widgets.return_value = None
        page.get_drawings.return_value = []
        document = MagicMock()
        document.__len__.return_value = 1
        document.__getitem__.return_value = page

        with patch("pdf_fidelity.pymupdf.open", return_value=document):
            result = analyze_pdf_fidelity_risk("partial-table-extraction.pdf")

        self.assertEqual(result.table_text_extraction_failure_count, 1)
        self.assertEqual(result.table_texts, ("Alpha Beta Gamma",))
        self.assertEqual(result.table_text_character_count, 14)
        self.assertEqual(result.table_shapes, ((2, 2),))
        self.assertEqual(result.table_cell_counts, (2,))
        self.assertTrue(result.table_analysis_uncertain)

    def test_docx_semantic_table_is_reported_as_editable(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "editable-table.docx"
            document = Document()
            document.add_paragraph("Document title")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Name"
            table.cell(0, 1).text = "Score"
            table.cell(1, 0).text = "Anonymous"
            table.cell(1, 1).text = "95"
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertTrue(summary.has_editable_table)
        self.assertEqual(summary.table_count, 1)
        self.assertEqual(summary.row_count, 2)
        self.assertEqual(summary.cell_count, 4)
        self.assertEqual(summary.table_shapes, ((2, 2),))
        self.assertEqual(summary.table_cell_counts, (4,))
        self.assertGreater(summary.table_text_character_count, 0)
        self.assertEqual(summary.table_texts, ("NameScoreAnonymous95",))
        self.assertEqual(
            summary.table_cell_matrices,
            ((("Name", "Score"), ("Anonymous", "95")),),
        )
        expected_visible_text = "Document titleNameScoreAnonymous95"
        self.assertEqual(summary.visible_text, expected_visible_text)
        self.assertEqual(
            summary.visible_text_character_count,
            sum(not character.isspace() for character in expected_visible_text),
        )
        self.assertFalse(summary.document_protected)

    def test_docx_explicit_cell_border_edges_preserve_sorted_span_order(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "explicit-borders.docx"
            document = Document()
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "A"
            table.cell(0, 1).text = "B"
            self._set_explicit_cell_borders(
                table.cell(0, 0),
                top="nil",
                left="single",
                bottom="single",
                right="nil",
            )
            self._set_explicit_cell_borders(
                table.cell(0, 1),
                top="single",
                left="single",
                bottom="single",
                right="single",
            )
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(
            summary.table_cell_spans,
            (((0, 0, 1, 1), (0, 1, 1, 1)),),
        )
        self.assertEqual(
            summary.table_cell_border_edges,
            (((False, True, True, False), (True, True, True, True)),),
        )

    def test_docx_table_style_borders_are_not_reported_as_explicit_edges(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "style-only-borders.docx"
            document = Document()
            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            table.cell(0, 0).text = "A"
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(
            summary.table_cell_border_edges,
            (((False, False, False, False),),),
        )

    def test_docx_vertical_merge_uses_only_logical_outer_border_fragments(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "merged-borders.docx"
            document = Document()
            table = document.add_table(rows=3, cols=1)
            table.cell(0, 0).merge(table.cell(2, 0)).text = "Merged"
            physical_cells = tuple(
                row.findall(qn("w:tc"))[0]
                for row in table._tbl.findall(qn("w:tr"))
            )
            self._set_explicit_cell_borders(
                physical_cells[0],
                top="single",
                left="single",
                bottom="single",
                right="single",
            )
            self._set_explicit_cell_borders(
                physical_cells[1],
                top="single",
                left="single",
                bottom="single",
                right="nil",
            )
            self._set_explicit_cell_borders(
                physical_cells[2],
                top="single",
                left="single",
                bottom="single",
                right="single",
            )
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.table_cell_spans, (((0, 0, 3, 1),),))
        self.assertEqual(
            summary.table_cell_border_edges,
            (((True, True, True, False),),),
        )

    def test_pdf_and_docx_cell_matrices_pass_end_to_end_gate(self):
        from pdf_word_converter import ConverterApp

        with tempfile.TemporaryDirectory() as temp_dir:
            folder = Path(temp_dir)
            source = folder / "matrix-source.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                self._draw_table(page)
                document.save(source)
            finally:
                document.close()

            report = analyze_pdf_fidelity_risk(
                source,
                max_pages=None,
            )
            output = folder / "matrix-output.docx"
            word_document = Document()
            table = word_document.add_table(rows=3, cols=3)
            for row in range(3):
                for column in range(3):
                    table.cell(row, column).text = (
                        f"R{row + 1}C{column + 1}"
                    )
            for cell in table._cells:
                self._set_explicit_cell_borders(
                    cell,
                    top="single",
                    left="single",
                    bottom="single",
                    right="single",
                )
            word_document.save(output)

            progress = MagicMock()
            ConverterApp._validate_editable_table_output(
                output,
                report,
                progress,
            )
            progress.assert_called_once()

            table.cell(1, 1).text = "WRONG"
            word_document.save(output)
            with self.assertRaisesRegex(RuntimeError, "0/1"):
                ConverterApp._validate_editable_table_output(
                    output,
                    report,
                    MagicMock(),
                )

            table.cell(1, 1).text = "R2C2"
            positioning = OxmlElement("w:tblpPr")
            positioning.set(qn("w:tblpX"), "100")
            positioning.set(qn("w:tblpY"), "100")
            table._tbl.tblPr.append(positioning)
            word_document.save(output)
            with self.assertRaisesRegex(RuntimeError, "清晰可见"):
                ConverterApp._validate_editable_table_output(
                    output,
                    report,
                    MagicMock(),
                )

    def test_docx_table_with_hidden_text_is_not_editable_output(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "hidden-table-text.docx"
            document = Document()
            table = document.add_table(rows=2, cols=2)
            values = ("Name", "Score", "Anonymous", "95")
            for cell, value in zip(table._cells, values):
                cell.text = value
                for run in cell.paragraphs[0].runs:
                    run._r.get_or_add_rPr().append(
                        OxmlElement("w:vanish")
                    )
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.table_count, 1)
        self.assertEqual(summary.table_text_character_count, 0)
        self.assertFalse(summary.has_editable_table)
    def test_docx_table_text_hidden_by_inherited_character_style_is_rejected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "style-hidden-table.docx"
            document = Document()
            base_style = document.styles.add_style(
                "InvisibleCellBase",
                WD_STYLE_TYPE.CHARACTER,
            )
            base_style.font.size = Pt(1)
            base_style.font.color.rgb = RGBColor(255, 255, 255)
            inherited_style = document.styles.add_style(
                "InvisibleCellInherited",
                WD_STYLE_TYPE.CHARACTER,
            )
            inherited_style.base_style = base_style
            table = document.add_table(rows=2, cols=2)
            for cell, value in zip(table._cells, ("A", "B", "C", "D")):
                cell.text = ""
                run = cell.paragraphs[0].add_run(value)
                run.style = inherited_style
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.table_count, 1)
        self.assertEqual(summary.table_text_character_count, 0)
        self.assertFalse(summary.has_editable_table)

    def test_text_boxes_inside_empty_cells_do_not_count_as_table_text(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "textbox-only-table.docx"
            document = Document()
            table = document.add_table(rows=1, cols=2)
            for cell, value in zip(table._cells, ("A", "B")):
                cell.text = ""
                run = OxmlElement("w:r")
                picture = OxmlElement("w:pict")
                text_box_content = OxmlElement("w:txbxContent")
                paragraph = OxmlElement("w:p")
                inner_run = OxmlElement("w:r")
                text_node = OxmlElement("w:t")
                text_node.text = value
                inner_run.append(text_node)
                paragraph.append(inner_run)
                text_box_content.append(paragraph)
                picture.append(text_box_content)
                run.append(picture)
                cell.paragraphs[0]._p.append(run)
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.table_count, 1)
        self.assertEqual(summary.table_texts, ("",))
        self.assertEqual(summary.table_text_character_count, 0)
        self.assertEqual(summary.invalid_layout_table_count, 1)
        self.assertFalse(summary.has_editable_table_structure)

    def test_floating_table_positioned_outside_page_is_rejected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "off-page-table.docx"
            document = Document()
            table = document.add_table(rows=2, cols=2)
            for cell, value in zip(table._cells, ("A", "B", "C", "D")):
                cell.text = value
            positioning = OxmlElement("w:tblpPr")
            positioning.set(qn("w:tblpX"), "30000")
            positioning.set(qn("w:tblpY"), "30000")
            table._tbl.tblPr.append(positioning)
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.invalid_layout_table_count, 1)
        self.assertFalse(summary.has_editable_table_structure)
        self.assertFalse(summary.has_editable_table)

    def test_one_twip_exact_table_rows_are_rejected_as_clipped(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "clipped-table.docx"
            document = Document()
            table = document.add_table(rows=2, cols=2)
            for cell, value in zip(table._cells, ("A", "B", "C", "D")):
                cell.text = value
            for row in table.rows:
                height = OxmlElement("w:trHeight")
                height.set(qn("w:val"), "1")
                height.set(qn("w:hRule"), "exact")
                row._tr.get_or_add_trPr().append(height)
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.invalid_layout_table_count, 1)
        self.assertFalse(summary.has_editable_table_structure)
        self.assertFalse(summary.has_editable_table)

    def test_one_twip_table_grid_column_is_rejected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "zero-width-column.docx"
            document = Document()
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "A"
            table.cell(0, 1).text = "B"
            grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
            self.assertEqual(len(grid_columns), 2)
            grid_columns[0].set(qn("w:w"), "1")
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.invalid_layout_table_count, 1)
        self.assertFalse(summary.has_editable_table_structure)
    def test_standard_white_header_on_dark_table_style_remains_visible(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "styled-visible-table.docx"
            document = Document()
            table = document.add_table(rows=2, cols=2)
            table.style = "Medium Shading 1 Accent 1"
            for cell, value in zip(table._cells, ("H1", "H2", "A", "B")):
                cell.text = value
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.table_texts, ("H1H2AB",))
        self.assertEqual(summary.invalid_layout_table_count, 0)
        self.assertTrue(summary.has_editable_table)

    def test_direct_black_text_overrides_inherited_white_style(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "visible-style-override.docx"
            document = Document()
            white_style = document.styles.add_style(
                "WhiteCellText",
                WD_STYLE_TYPE.CHARACTER,
            )
            white_style.font.color.rgb = RGBColor(255, 255, 255)
            table = document.add_table(rows=1, cols=2)
            for cell, value in zip(table._cells, ("Name", "Score")):
                cell.text = ""
                run = cell.paragraphs[0].add_run(value)
                run.style = white_style
                run.font.color.rgb = RGBColor(0, 0, 0)
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.table_texts, ("NameScore",))
        self.assertEqual(summary.invalid_layout_table_count, 0)
        self.assertTrue(summary.has_editable_table)

    def test_white_table_text_on_light_gray_is_rejected_for_low_contrast(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "low-contrast-table.docx"
            document = Document()
            table = document.add_table(rows=1, cols=2)
            for cell, value in zip(table._cells, ("Name", "Score")):
                cell.text = value
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "C0C0C0")
                cell._tc.get_or_add_tcPr().append(shading)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.table_text_character_count, 0)
        self.assertEqual(summary.invalid_layout_table_count, 1)
        self.assertFalse(summary.has_editable_table)

    def test_any_floating_table_is_rejected_when_visibility_cannot_be_proven(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "floating-table.docx"
            document = Document()
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "A"
            table.cell(0, 1).text = "B"
            positioning = OxmlElement("w:tblpPr")
            positioning.set(qn("w:tblpX"), "100")
            positioning.set(qn("w:tblpY"), "100")
            table._tbl.tblPr.append(positioning)
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.invalid_layout_table_count, 1)
        self.assertFalse(summary.has_editable_table_structure)

    def test_extremely_small_percentage_table_width_is_rejected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "tiny-percentage-table.docx"
            document = Document()
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "A"
            table.cell(0, 1).text = "B"
            table_width = table._tbl.tblPr.find(qn("w:tblW"))
            self.assertIsNotNone(table_width)
            table_width.set(qn("w:type"), "pct")
            table_width.set(qn("w:w"), "6")
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.invalid_layout_table_count, 1)
        self.assertFalse(summary.has_editable_table_structure)

    def test_missing_grid_column_width_is_rejected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "missing-grid-width.docx"
            document = Document()
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "A"
            table.cell(0, 1).text = "B"
            grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
            self.assertEqual(len(grid_columns), 2)
            grid_columns[0].attrib.pop(qn("w:w"), None)
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.invalid_layout_table_count, 1)
        self.assertFalse(summary.has_editable_table_structure)

    def test_text_outside_normal_cell_paragraphs_invalidates_table(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "malformed-table-text.docx"
            document = Document()
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Visible"
            bogus_text = OxmlElement("w:t")
            bogus_text.text = "Hidden outside a cell"
            table._tbl.tblPr.append(bogus_text)
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.invalid_layout_table_count, 1)
        self.assertFalse(summary.has_editable_table_structure)
    def test_docx_grid_span_counts_effective_columns(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "merged-table.docx"
            document = Document()
            table = document.add_table(rows=2, cols=3)
            table.cell(0, 0).merge(table.cell(0, 1)).text = "Merged heading"
            table.cell(0, 2).text = "Result"
            for column, value in enumerate(("A", "B", "C")):
                table.cell(1, column).text = value
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.table_shapes, ((2, 3),))
        self.assertEqual(summary.table_cell_counts, (5,))
        self.assertEqual(
            summary.table_cell_matrices,
            (
                (
                    ("Merged heading", None, "Result"),
                    ("A", "B", "C"),
                ),
            ),
        )
        self.assertTrue(summary.has_editable_table)
    def test_docx_vertical_merge_counts_one_independent_cell(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "vertical-merge.docx"
            document = Document()
            table = document.add_table(rows=10, cols=2)
            table.cell(0, 0).merge(table.cell(9, 0)).text = "Merged group"
            for row in range(10):
                table.cell(row, 1).text = f"Value {row + 1}"
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.table_shapes, ((10, 2),))
        self.assertEqual(summary.table_cell_counts, (11,))
        self.assertTrue(summary.has_editable_table)
    def test_docx_full_row_gridspan_reports_physical_cell_collapse(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "full-row-merge.docx"
            document = Document()
            table = document.add_table(rows=3, cols=4)
            for row in range(3):
                table.cell(row, 0).merge(table.cell(row, 3)).text = f"Row {row + 1}"
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.table_shapes, ((3, 4),))
        self.assertEqual(summary.table_cell_counts, (3,))
        self.assertTrue(summary.has_editable_table)
    def test_nested_table_text_is_counted_once(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "nested-table.docx"
            document = Document()
            outer_table = document.add_table(rows=1, cols=1)
            inner_table = outer_table.cell(0, 0).add_table(rows=1, cols=1)
            inner_table.cell(0, 0).text = "X" * 40
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertTrue(summary.has_editable_table)
        self.assertEqual(summary.table_count, 2)
        self.assertEqual(summary.table_text_character_count, 40)
        self.assertEqual(summary.table_texts, ("", "X" * 40))
        self.assertEqual(
            summary.table_cell_matrices,
            ((("",),), (("X" * 40,),)),
        )

    def test_docx_paragraph_and_empty_table_are_not_editable_table_results(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            folder = Path(temp_dir)
            paragraph_output = folder / "paragraph-only.docx"
            paragraph_document = Document()
            paragraph_document.add_paragraph("Table-looking text")
            paragraph_document.save(paragraph_output)

            empty_output = folder / "empty-table.docx"
            empty_document = Document()
            empty_document.add_table(rows=1, cols=1)
            empty_document.save(empty_output)

            paragraph_summary = inspect_editable_docx_tables(paragraph_output)
            empty_summary = inspect_editable_docx_tables(empty_output)

        self.assertFalse(paragraph_summary.has_editable_table)
        self.assertEqual(paragraph_summary.table_count, 0)
        self.assertFalse(empty_summary.has_editable_table)
        self.assertTrue(empty_summary.has_editable_table_structure)
        self.assertEqual(empty_summary.table_count, 1)
        self.assertEqual(empty_summary.table_text_character_count, 0)

    def test_image_only_docx_is_not_an_editable_table_result(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            folder = Path(temp_dir)
            image_path = folder / "page.png"
            output = folder / "image-only.docx"
            Image.new("RGB", (100, 100), "white").save(image_path)
            document = Document()
            document.add_picture(str(image_path))
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertFalse(summary.has_editable_table)
        self.assertEqual(summary.table_count, 0)
        self.assertEqual(summary.table_text_character_count, 0)
        self.assertEqual(summary.drawing_count, 1)

    def test_protected_docx_table_is_rejected_as_editable_output(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "protected-table.docx"
            document = Document()
            document.add_table(rows=1, cols=1).cell(0, 0).text = "Locked"
            protection = OxmlElement("w:documentProtection")
            protection.set(qn("w:edit"), "readOnly")
            protection.set(qn("w:enforcement"), "1")
            document.settings.element.append(protection)
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertTrue(summary.document_protected)
        self.assertFalse(summary.has_editable_table)

    def test_invalid_docx_is_reported_clearly(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "broken.docx"
            output.write_bytes(b"not a docx")
            with self.assertRaisesRegex(RuntimeError, "Word"):
                inspect_editable_docx_tables(output)

    def test_fake_docx_zip_without_package_parts_is_rejected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "fake.docx"
            with ZipFile(output, "w") as archive:
                archive.writestr(
                    "word/document.xml",
                    (
                        '<?xml version="1.0" encoding="UTF-8"?>'
                        '<w:document xmlns:w="'
                        'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                        '"><w:body><w:tbl><w:tr><w:tc><w:p><w:r><w:t>'
                        + ("X" * 80)
                        + "</w:t></w:r></w:p></w:tc></w:tr></w:tbl>"
                        "</w:body></w:document>"
                    ),
                )

            with self.assertRaisesRegex(RuntimeError, "无法打开 Word 文档"):
                inspect_editable_docx_tables(output)

    def test_locked_content_control_table_is_not_editable(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "locked-content-control.docx"
            document = Document()
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Locked cell"
            table_element = table._tbl
            parent = table_element.getparent()
            insertion_index = parent.index(table_element)
            parent.remove(table_element)

            content_control = OxmlElement("w:sdt")
            properties = OxmlElement("w:sdtPr")
            lock = OxmlElement("w:lock")
            lock.set(qn("w:val"), "sdtContentLocked")
            properties.append(lock)
            content = OxmlElement("w:sdtContent")
            content.append(table_element)
            content_control.append(properties)
            content_control.append(content)
            parent.insert(insertion_index, content_control)
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.table_count, 1)
        self.assertEqual(summary.locked_content_control_table_count, 1)
        self.assertFalse(summary.has_editable_table)

    def test_locked_content_control_inside_table_is_not_editable(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "locked-cell-content-control.docx"
            document = Document()
            table = document.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            cell.text = ""

            content_control = OxmlElement("w:sdt")
            properties = OxmlElement("w:sdtPr")
            lock = OxmlElement("w:lock")
            lock.set(qn("w:val"), "contentLocked")
            properties.append(lock)
            content = OxmlElement("w:sdtContent")
            paragraph = OxmlElement("w:p")
            run = OxmlElement("w:r")
            text = OxmlElement("w:t")
            text.text = "Locked cell"
            run.append(text)
            paragraph.append(run)
            content.append(paragraph)
            content_control.append(properties)
            content_control.append(content)
            cell._tc.append(content_control)
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.table_count, 1)
        self.assertEqual(summary.locked_content_control_table_count, 1)
        self.assertFalse(summary.has_editable_table)

    def test_write_protected_docx_table_is_not_editable(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "write-protected-table.docx"
            document = Document()
            document.add_table(rows=1, cols=1).cell(0, 0).text = "Protected"
            protection = OxmlElement("w:writeProtection")
            document.settings.element.append(protection)
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertTrue(summary.document_protected)
        self.assertFalse(summary.has_editable_table)

    def test_invalid_limits_are_rejected(self):
        with self.assertRaisesRegex(ValueError, "max_pages"):
            analyze_pdf_fidelity_risk("unused.pdf", max_pages=0)
        with self.assertRaisesRegex(ValueError, "vector_threshold"):
            analyze_pdf_fidelity_risk("unused.pdf", vector_threshold=-1)

    def test_pdf_fidelity_analysis_can_cancel_before_opening_source(self):
        with patch("pdf_fidelity.pymupdf.open") as open_pdf:
            with self.assertRaises(PdfFidelityAnalysisCancelled):
                analyze_pdf_fidelity_risk(
                    "unused.pdf",
                    cancel_requested=lambda: True,
                )

        open_pdf.assert_not_called()

    def test_pdf_fidelity_analysis_cancels_between_text_table_strategies(self):
        cancellation_state = {"requested": False}
        find_table_calls = []

        def find_tables(**kwargs):
            find_table_calls.append(dict(kwargs))
            if (
                kwargs.get("strategy") == "text"
                and kwargs.get("min_words_vertical") == 2
            ):
                cancellation_state["requested"] = True
            finder = MagicMock()
            finder.tables = ()
            return finder

        page = MagicMock()
        page.rect = pymupdf.Rect(0, 0, 595, 842)
        page.rotation = 0
        page.find_tables.side_effect = find_tables
        page.get_text.side_effect = lambda mode, *args, **kwargs: (
            {"blocks": [{"lines": [{"spans": [{"text": "A", "font": "Arial"}]}]}]}
            if mode == "dict"
            else ()
        )
        page.get_drawings.return_value = ()
        page.widgets.return_value = None
        document = MagicMock()
        document.__len__.return_value = 1
        document.__getitem__.return_value = page

        with (
            patch("pdf_fidelity.pymupdf.open", return_value=document),
            patch(
                "pdf_fidelity._document_has_tagged_table_structure",
                return_value=False,
            ),
            patch("pdf_fidelity._page_has_table_like_layout", return_value=False),
        ):
            with self.assertRaises(PdfFidelityAnalysisCancelled):
                analyze_pdf_fidelity_risk(
                    "cancel-between-strategies.pdf",
                    cancel_requested=lambda: cancellation_state["requested"],
                )

        self.assertIn(
            {"strategy": "text", "min_words_vertical": 2},
            find_table_calls,
        )
        self.assertNotIn(
            {"strategy": "text", "min_words_vertical": 1},
            find_table_calls,
        )
        document.close.assert_called_once_with()

    def test_pdf_fidelity_analysis_rejects_non_callable_cancel_probe(self):
        with self.assertRaisesRegex(ValueError, "cancel_requested"):
            analyze_pdf_fidelity_risk("unused.pdf", cancel_requested=True)

    def test_blank_grid_without_extractable_cells_fails_closed(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "blank-editable-grid.pdf"
            document = pymupdf.open()
            try:
                page = document.new_page(width=595, height=842)
                left, top, cell_width, cell_height = 72, 90, 110, 42
                rows, columns = 3, 3
                for row in range(rows + 1):
                    y = top + row * cell_height
                    page.draw_line(
                        (left, y),
                        (left + columns * cell_width, y),
                        color=(0, 0, 0),
                    )
                for column in range(columns + 1):
                    x = left + column * cell_width
                    page.draw_line(
                        (x, top),
                        (x, top + rows * cell_height),
                        color=(0, 0, 0),
                    )
                document.save(source)
            finally:
                document.close()

            result = analyze_pdf_fidelity_risk(source, max_pages=None)

        self.assertEqual(result.table_count, 0)
        self.assertEqual(result.table_text_character_count, 0)
        self.assertFalse(result.editable_table_candidate)
        self.assertTrue(result.table_layout_suspected)

    def test_white_table_text_and_large_page_image_are_detected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            folder = Path(temp_dir)
            image_path = folder / "page.png"
            output = folder / "photo-with-white-table.docx"
            Image.new("RGB", (1200, 1600), "white").save(image_path)
            document = Document()
            table = document.add_table(rows=2, cols=2)
            for cell, value in zip(table._cells, ("A", "B", "C", "D")):
                cell.text = value
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            document.add_picture(
                str(image_path),
                width=Inches(7.2),
                height=Inches(9.0),
            )
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.table_count, 1)
        self.assertEqual(summary.table_text_character_count, 0)
        self.assertFalse(summary.has_editable_table)
        self.assertEqual(summary.large_page_drawing_count, 1)

    def test_white_table_text_on_dark_cells_remains_visible(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "visible-white-on-dark.docx"
            document = Document()
            table = document.add_table(rows=1, cols=2)
            for cell, value in zip(table._cells, ("Name", "Score")):
                cell.text = value
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "000000")
                cell._tc.get_or_add_tcPr().append(shading)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            document.save(output)

            summary = inspect_editable_docx_tables(output)

        self.assertEqual(summary.table_texts, ("NameScore",))
        self.assertTrue(summary.has_editable_table)
        self.assertEqual(summary.large_page_drawing_count, 0)

    def test_nested_source_table_candidate_is_not_removed_as_duplicate(self):
        outer = MagicMock()
        outer.bbox = (10, 10, 500, 500)
        outer.col_count = 2
        outer.extract.return_value = (("Outer", "Table"),)
        nested = MagicMock()
        nested.bbox = (100, 100, 300, 300)
        nested.col_count = 2
        nested.extract.return_value = (("Nested", "Values"),)
        union = MagicMock()
        union.bbox = (5, 5, 510, 510)
        union.col_count = 4
        union.extract.return_value = (("Outer", "Nested", "Table", "Values"),)

        result = _deduplicate_table_candidates((outer, nested, union))

        self.assertEqual(result, (outer, nested))


    def test_precise_grid_rejects_unowned_interleaved_narrow_column(self):
        default = object()
        precise = object()
        default_rows = (("A", "B"), ("C", "D"))
        precise_rows = (("A", None, "B"), ("C", None, "D"))
        shapes = {default: (2, 2), precise: (2, 3)}
        geometries = {
            default: (
                (
                    (0, 0, 1, 1),
                    (0, 1, 1, 1),
                    (1, 0, 1, 1),
                    (1, 1, 1, 1),
                ),
                (50.0, 50.0),
            ),
            precise: (
                (
                    (0, 0, 1, 2),
                    (0, 2, 1, 1),
                    (1, 0, 1, 2),
                    (1, 2, 1, 1),
                ),
                (49.0, 2.0, 49.0),
            ),
        }

        with (
            patch(
                "pdf_fidelity._table_rect",
                return_value=pymupdf.Rect(0.0, 0.0, 100.0, 20.0),
            ),
            patch(
                "pdf_fidelity._effective_table_shape",
                side_effect=lambda table, _rows: shapes[table],
            ),
            patch(
                "pdf_fidelity._effective_table_cell_count",
                return_value=4,
            ),
            patch(
                "pdf_fidelity._table_geometry_model",
                side_effect=lambda table, _shape: geometries[table],
            ),
        ):
            selected = _select_precise_line_tables(
                (default,),
                (precise,),
                {id(default): default_rows, id(precise): precise_rows},
            )

        self.assertIs(selected[0], default)

    def test_precise_grid_keeps_narrow_column_with_unit_span_owner(self):
        default = object()
        precise = object()
        default_rows = (("A", "B"), ("C", "D"))
        precise_rows = (("A", None, "B"), ("C", None, "D"))
        shapes = {default: (2, 2), precise: (2, 3)}
        geometries = {
            default: (
                (
                    (0, 0, 1, 1),
                    (0, 1, 1, 1),
                    (1, 0, 1, 1),
                    (1, 1, 1, 1),
                ),
                (50.0, 50.0),
            ),
            precise: (
                (
                    (0, 0, 1, 1),
                    (0, 1, 1, 1),
                    (0, 2, 2, 1),
                    (1, 0, 1, 2),
                ),
                (49.0, 2.0, 49.0),
            ),
        }

        with (
            patch(
                "pdf_fidelity._table_rect",
                return_value=pymupdf.Rect(0.0, 0.0, 100.0, 20.0),
            ),
            patch(
                "pdf_fidelity._effective_table_shape",
                side_effect=lambda table, _rows: shapes[table],
            ),
            patch(
                "pdf_fidelity._effective_table_cell_count",
                return_value=4,
            ),
            patch(
                "pdf_fidelity._table_geometry_model",
                side_effect=lambda table, _shape: geometries[table],
            ),
        ):
            selected = _select_precise_line_tables(
                (default,),
                (precise,),
                {id(default): default_rows, id(precise): precise_rows},
            )

        self.assertIs(selected[0], precise)

    def test_precise_only_table_candidate_is_retained(self):
        precise = MagicMock()
        precise.bbox = (220.0, 300.0, 420.0, 420.0)
        precise.col_count = 2
        precise_rows = (("P1", "P2"),)

        selected = _select_precise_line_tables(
            (),
            (precise,),
            {id(precise): precise_rows},
        )

        self.assertEqual(selected, (precise,))

    def test_detected_table_does_not_hide_unconfirmed_residual_region(self):
        table = MagicMock()
        table.bbox = (10.0, 20.0, 110.0, 60.0)
        table.row_count = 1
        table.col_count = 1
        table.cells = (table.bbox,)
        table.extract.return_value = (("A",),)
        row = MagicMock()
        row.bbox = table.bbox
        row.cells = (table.bbox,)
        table.rows = (row,)

        page = MagicMock()
        page.rect = pymupdf.Rect(0, 0, 595, 842)
        page.find_tables.return_value.tables = (table,)
        page.get_text.return_value = {"blocks": []}
        page.get_textbox.return_value = "A"
        page.get_drawings.return_value = ()
        page.widgets.return_value = None

        document = MagicMock()
        document.__len__.return_value = 1
        document.__getitem__.return_value = page

        with (
            patch("pdf_fidelity.pymupdf.open", return_value=document),
            patch(
                "pdf_fidelity._document_has_tagged_table_structure",
                return_value=False,
            ),
            patch(
                "pdf_fidelity._page_has_table_like_layout",
                side_effect=(True, True),
            ) as layout_probe,
        ):
            result = analyze_pdf_fidelity_risk("partial-evidence.pdf")

        self.assertEqual(result.table_count, 1)
        self.assertEqual(result.unconfirmed_table_region_count, 1)
        self.assertTrue(any("无法确认的表格区域" in reason for reason in result.reasons))
        self.assertEqual(layout_probe.call_count, 2)
        self.assertEqual(
            layout_probe.call_args_list[1].kwargs["excluded_rects"],
            (pymupdf.Rect(table.bbox),),
        )
        document.close.assert_called_once_with()

    def test_rotated_fully_detected_table_has_no_unconfirmed_residual(self):
        table = MagicMock()
        table.bbox = (16.0, 71.0, 826.0, 544.0)
        table.row_count = 1
        table.col_count = 1
        table.cells = (table.bbox,)
        table.extract.return_value = (("A",),)
        row = MagicMock()
        row.bbox = table.bbox
        row.cells = (table.bbox,)
        table.rows = (row,)

        page = MagicMock()
        page.rect = pymupdf.Rect(0, 0, 842, 595)
        page.rotation = 90
        page.derotation_matrix = pymupdf.Matrix(0, -1, 1, 0, 0, 842)
        page.find_tables.return_value.tables = (table,)
        page.get_text.return_value = {"blocks": []}
        page.get_textbox.return_value = "A"
        page.get_drawings.return_value = ()
        page.widgets.return_value = None

        document = MagicMock()
        document.__len__.return_value = 1
        document.__getitem__.return_value = page

        with (
            patch("pdf_fidelity.pymupdf.open", return_value=document),
            patch(
                "pdf_fidelity._document_has_tagged_table_structure",
                return_value=False,
            ),
            patch(
                "pdf_fidelity._page_has_table_like_layout",
                side_effect=(True, False),
            ) as layout_probe,
        ):
            result = analyze_pdf_fidelity_risk("rotated-full-table.pdf")

        self.assertEqual(result.table_count, 1)
        self.assertEqual(result.unconfirmed_table_region_count, 0)
        self.assertEqual(layout_probe.call_count, 2)
        self.assertEqual(
            layout_probe.call_args_list[1].kwargs["excluded_rects"],
            (pymupdf.Rect(71.0, 16.0, 544.0, 826.0),),
        )
        document.close.assert_called_once_with()

if __name__ == "__main__":
    unittest.main()
