import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips
from lxml import etree

from docx_table_repair import DocxTableRepairError, repair_docx_table_topology
from pdf_fidelity import inspect_editable_docx_tables


class DocxTableRepairTests(unittest.TestCase):
    @staticmethod
    def _replace_grid_widths(table, widths):
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            grid_column = OxmlElement("w:gridCol")
            grid_column.set(qn("w:w"), str(int(width)))
            grid.append(grid_column)

    @staticmethod
    def _set_cell_width(cell, width):
        properties = cell._tc.get_or_add_tcPr()
        cell_width = properties.find(qn("w:tcW"))
        if cell_width is None:
            cell_width = OxmlElement("w:tcW")
            properties.insert(0, cell_width)
        cell_width.set(qn("w:type"), "dxa")
        cell_width.set(qn("w:w"), str(int(width)))

    @staticmethod
    def _set_cell_borders(
        cell,
        *,
        include_top=True,
        nil_horizontal=False,
    ):
        properties = cell._tc.get_or_add_tcPr()
        borders = properties.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            properties.append(borders)
        for child in list(borders):
            borders.remove(child)
        edges = ("top", "left", "bottom", "right")
        for edge_name in edges:
            if edge_name == "top" and not include_top:
                continue
            edge = OxmlElement(f"w:{edge_name}")
            edge.set(
                qn("w:val"),
                (
                    "nil"
                    if nil_horizontal and edge_name in {"top", "bottom"}
                    else "single"
                ),
            )
            edge.set(qn("w:sz"), "4")
            edge.set(qn("w:space"), "0")
            edge.set(qn("w:color"), "auto")
            borders.append(edge)

    @staticmethod
    def _set_run_border(run):
        properties = run._r.get_or_add_rPr()
        border = OxmlElement("w:bdr")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        properties.append(border)

    @classmethod
    def _create_table_document(
        cls,
        path,
        values,
        *,
        grid_widths=None,
        row_widths=None,
    ):
        document = Document()
        table = document.add_table(rows=len(values), cols=len(values[0]))
        table.autofit = False
        if grid_widths is None:
            grid_widths = (1000,) * len(values[0])
        cls._replace_grid_widths(table, grid_widths)
        for row_index, row_values in enumerate(values):
            widths = row_widths[row_index] if row_widths else grid_widths
            for column_index, value in enumerate(row_values):
                cell = table.cell(row_index, column_index)
                cell.text = value
                cls._set_cell_width(cell, widths[column_index])
        document.save(path)
        return document, table

    @staticmethod
    def _repair_one(
        path,
        *,
        shape,
        matrix,
        spans,
        widths,
        border_edges=None,
    ):
        return repair_docx_table_topology(
            path,
            table_shapes=(shape,),
            table_cell_matrices=(matrix,),
            table_cell_spans=(spans,),
            table_column_widths=(widths,),
            table_cell_border_edges=(
                (border_edges,) if border_edges is not None else None
            ),
        )

    @staticmethod
    def _unit_spans(rows, columns):
        return tuple(
            (row, column, 1, 1)
            for row in range(rows)
            for column in range(columns)
        )

    @staticmethod
    def _row_grid_spans(table):
        result = []
        for row in table._tbl.findall(qn("w:tr")):
            row_spans = []
            for cell in row.findall(qn("w:tc")):
                grid_span = cell.find("./w:tcPr/w:gridSpan", cell.nsmap)
                row_spans.append(
                    int(grid_span.get(qn("w:val"), "1"))
                    if grid_span is not None
                    else 1
                )
            result.append(tuple(row_spans))
        return tuple(result)

    @staticmethod
    def _raw_cell_at(table, row_index, column_index):
        row = table._tbl.findall(qn("w:tr"))[row_index]
        current_column = 0
        for cell in row.findall(qn("w:tc")):
            grid_span = cell.find("./w:tcPr/w:gridSpan", cell.nsmap)
            span = (
                int(grid_span.get(qn("w:val"), "1"))
                if grid_span is not None
                else 1
            )
            if current_column == column_index:
                return cell
            current_column += span
        raise AssertionError(f"No raw cell starts at column {column_index}")

    @classmethod
    def _raw_cell_border_values(cls, table, row_index, column_index):
        cell = cls._raw_cell_at(table, row_index, column_index)
        values = []
        for name in ("top", "left", "bottom", "right"):
            border = cell.find(f"./w:tcPr/w:tcBorders/w:{name}", cell.nsmap)
            values.append(
                border.get(qn("w:val")) if border is not None else None
            )
        return tuple(values)

    @staticmethod
    def _assert_unchanged_after_error(test_case, path, call):
        original_bytes = path.read_bytes()
        with test_case.assertRaises(DocxTableRepairError):
            call()
        test_case.assertEqual(path.read_bytes(), original_bytes)

    @classmethod
    def _create_aggregated_vertical_merge(cls, path, *, split_runs):
        document, table = cls._create_table_document(
            path,
            (("R1", "", ""), ("R2", "", "")),
        )
        merged = table.cell(0, 1).merge(table.cell(1, 2))
        merged.text = ""
        first_paragraph = merged.paragraphs[0]
        second_paragraph = merged.add_paragraph()
        if split_runs:
            first_paragraph.add_run("1.5").bold = True
            first_paragraph.add_run("81").italic = True
            second_paragraph.add_run("3.0").underline = True
            second_paragraph.add_run("86").font.size = Pt(12)
        else:
            first_paragraph.add_run("1.581")
            second_paragraph.add_run("3.086")
        document.save(path)

    @classmethod
    def _create_cleanup_candidate(
        cls,
        path,
        *,
        border_count=0,
        incomplete_border_row=None,
        row_height=None,
        total_width=10000,
        nil_horizontal_borders=False,
    ):
        values = tuple((f"R{index + 1}", "") for index in range(4))
        first_width = total_width // 2
        grid_widths = (first_width, total_width - first_width)
        document, table = cls._create_table_document(
            path,
            values,
            grid_widths=grid_widths,
        )
        section = document.sections[0]
        section.page_width = Twips(12240)
        section.page_height = Twips(15840)
        section.top_margin = Twips(1440)
        section.bottom_margin = Twips(1440)
        section.left_margin = Twips(1800)
        section.right_margin = Twips(1800)

        for row_index, row in enumerate(table.rows):
            include_top = row_index != incomplete_border_row
            for cell in row.cells:
                cls._set_cell_borders(
                    cell,
                    include_top=include_top,
                    nil_horizontal=nil_horizontal_borders,
                )
            if row_index < border_count:
                cls._set_run_border(row.cells[0].paragraphs[0].runs[0])
            if row_height is not None:
                row.height = Twips(row_height)
        document.save(path)
        matrix = tuple((f"R{index + 1}",) for index in range(4))
        spans = cls._unit_spans(4, 1)
        return matrix, spans, (float(total_width),)

    def test_matching_topology_returns_false_without_rewriting_file(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "already-matched.docx"
            matrix = (("A", "B"), ("C", "D"))
            self._create_table_document(path, matrix)
            original_bytes = path.read_bytes()

            changed = self._repair_one(
                path,
                shape=(2, 2),
                matrix=matrix,
                spans=self._unit_spans(2, 2),
                widths=(100.0, 100.0),
            )

            self.assertFalse(changed)
            self.assertEqual(path.read_bytes(), original_bytes)

    def test_matching_topology_repairs_nil_borders_and_is_idempotent(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "matching-nil-borders.docx"
            document, table = self._create_table_document(
                path,
                (("A",), ("B",)),
            )
            for row in table.rows:
                self._set_cell_borders(
                    row.cells[0],
                    nil_horizontal=True,
                )
            document.save(path)
            matrix = (("A",), ("B",))
            spans = self._unit_spans(2, 1)
            border_edges = ((True, True, True, True),) * 2

            changed = self._repair_one(
                path,
                shape=(2, 1),
                matrix=matrix,
                spans=spans,
                widths=(100.0,),
                border_edges=border_edges,
            )

            self.assertTrue(changed)
            repaired = Document(path).tables[0]
            self.assertEqual(
                self._raw_cell_border_values(repaired, 0, 0),
                ("single", "single", "single", "single"),
            )
            self.assertEqual(
                self._raw_cell_border_values(repaired, 1, 0),
                ("single", "single", "single", "single"),
            )
            repaired_bytes = path.read_bytes()
            self.assertFalse(
                self._repair_one(
                    path,
                    shape=(2, 1),
                    matrix=matrix,
                    spans=spans,
                    widths=(100.0,),
                    border_edges=border_edges,
                )
            )
            self.assertEqual(path.read_bytes(), repaired_bytes)

    def test_explicit_borderless_evidence_writes_nil_without_adding_lines(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "borderless.docx"
            matrix = (("A",), ("B",))
            self._create_table_document(path, matrix)
            spans = self._unit_spans(2, 1)

            changed = self._repair_one(
                path,
                shape=(2, 1),
                matrix=matrix,
                spans=spans,
                widths=(100.0,),
                border_edges=((False, False, False, False),) * 2,
            )

            self.assertTrue(changed)
            repaired = Document(path).tables[0]
            self.assertEqual(
                self._raw_cell_border_values(repaired, 0, 0),
                ("nil", "nil", "nil", "nil"),
            )
            self.assertEqual(
                self._raw_cell_border_values(repaired, 1, 0),
                ("nil", "nil", "nil", "nil"),
            )

    def test_vertical_merge_keeps_only_outer_horizontal_borders(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "merged-border-chain.docx"
            document, table = self._create_table_document(
                path,
                (("Group", "A"), ("", "B"), ("", "C")),
            )
            for row in table.rows:
                for cell in row.cells:
                    self._set_cell_borders(cell)
            merged = table.cell(0, 0).merge(table.cell(2, 0))
            merged.text = "Group"
            document.save(path)
            matrix = (("Group", "A"), (None, "B"), (None, "C"))
            spans = (
                (0, 0, 3, 1),
                (0, 1, 1, 1),
                (1, 1, 1, 1),
                (2, 1, 1, 1),
            )
            border_edges = ((True, True, True, True),) * len(spans)

            changed = self._repair_one(
                path,
                shape=(3, 2),
                matrix=matrix,
                spans=spans,
                widths=(100.0, 100.0),
                border_edges=border_edges,
            )

            self.assertTrue(changed)
            repaired = Document(path).tables[0]
            self.assertEqual(
                self._raw_cell_border_values(repaired, 0, 0),
                ("single", "single", "nil", "single"),
            )
            self.assertEqual(
                self._raw_cell_border_values(repaired, 1, 0),
                ("nil", "single", "nil", "single"),
            )
            self.assertEqual(
                self._raw_cell_border_values(repaired, 2, 0),
                ("nil", "single", "single", "single"),
            )

    def test_existing_visible_border_style_is_preserved_and_reused(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "styled-border.docx"
            document, table = self._create_table_document(path, (("A",),))
            cell = table.cell(0, 0)
            properties = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            top = OxmlElement("w:top")
            top.set(qn("w:val"), "double")
            top.set(qn("w:sz"), "12")
            top.set(qn("w:space"), "0")
            top.set(qn("w:color"), "FF0000")
            borders.append(top)
            properties.append(borders)
            document.save(path)

            self.assertTrue(
                self._repair_one(
                    path,
                    shape=(1, 1),
                    matrix=(("A",),),
                    spans=((0, 0, 1, 1),),
                    widths=(100.0,),
                    border_edges=((True, True, True, True),),
                )
            )

            repaired = Document(path).tables[0]
            raw_cell = self._raw_cell_at(repaired, 0, 0)
            for name in ("top", "left", "bottom", "right"):
                border = raw_cell.find(
                    f"./w:tcPr/w:tcBorders/w:{name}",
                    raw_cell.nsmap,
                )
                self.assertIsNotNone(border)
                self.assertEqual(border.get(qn("w:val")), "double")
                self.assertEqual(border.get(qn("w:sz")), "12")
                self.assertEqual(border.get(qn("w:color")), "FF0000")

    def test_visible_border_fallback_is_explicit_black_single(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "fallback-border.docx"
            self._create_table_document(path, (("A",),))

            self.assertTrue(
                self._repair_one(
                    path,
                    shape=(1, 1),
                    matrix=(("A",),),
                    spans=((0, 0, 1, 1),),
                    widths=(100.0,),
                    border_edges=((True, True, True, True),),
                )
            )

            repaired = Document(path).tables[0]
            raw_cell = self._raw_cell_at(repaired, 0, 0)
            for name in ("top", "left", "bottom", "right"):
                border = raw_cell.find(
                    f"./w:tcPr/w:tcBorders/w:{name}",
                    raw_cell.nsmap,
                )
                self.assertIsNotNone(border)
                self.assertEqual(border.get(qn("w:val")), "single")
                self.assertEqual(border.get(qn("w:sz")), "4")
                self.assertEqual(border.get(qn("w:space")), "0")
                self.assertEqual(border.get(qn("w:color")), "000000")

    def test_complex_cell_border_is_rejected_atomically(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "diagonal-border.docx"
            document, table = self._create_table_document(path, (("A",),))
            properties = table.cell(0, 0)._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            diagonal = OxmlElement("w:tl2br")
            diagonal.set(qn("w:val"), "single")
            borders.append(diagonal)
            properties.append(borders)
            document.save(path)

            self._assert_unchanged_after_error(
                self,
                path,
                lambda: self._repair_one(
                    path,
                    shape=(1, 1),
                    matrix=(("A",),),
                    spans=((0, 0, 1, 1),),
                    widths=(100.0,),
                    border_edges=((True, True, True, True),),
                ),
            )

    def test_invalid_border_evidence_is_rejected_atomically(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "invalid-border-evidence.docx"
            matrix = (("A",), ("B",))
            self._create_table_document(path, matrix)
            spans = self._unit_spans(2, 1)
            invalid_cases = (
                ("wrong-cell-count", ((True, True, True, True),)),
                (
                    "non-boolean-edge",
                    ((True, True, True, True), (True, True, True, 1)),
                ),
                (
                    "conflicting-shared-edge",
                    ((True, True, True, True), (False, True, True, True)),
                ),
            )
            for name, border_edges in invalid_cases:
                with self.subTest(name=name):
                    self._assert_unchanged_after_error(
                        self,
                        path,
                        lambda border_edges=border_edges: self._repair_one(
                            path,
                            shape=(2, 1),
                            matrix=matrix,
                            spans=spans,
                            widths=(100.0,),
                            border_edges=border_edges,
                        ),
                    )

    def test_rebuilds_three_column_grid_from_different_near_row_boundaries(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "row-boundaries.docx"
            self._create_table_document(
                path,
                (("A", "B"), ("C", "D")),
                grid_widths=(1500, 1500),
                row_widths=((980, 2020), (1980, 1020)),
            )
            matrix = (
                ("A", "B", None),
                ("C", None, "D"),
            )
            spans = (
                (0, 0, 1, 1),
                (0, 1, 1, 2),
                (1, 0, 1, 2),
                (1, 2, 1, 1),
            )

            changed = self._repair_one(
                path,
                shape=(2, 3),
                matrix=matrix,
                spans=spans,
                widths=(100.0, 100.0, 100.0),
            )

            self.assertTrue(changed)
            summary = inspect_editable_docx_tables(path)
            self.assertEqual(summary.table_shapes, ((2, 3),))
            self.assertEqual(summary.table_cell_counts, (4,))
            self.assertEqual(summary.table_cell_matrices, (matrix,))
            self.assertEqual(summary.table_cell_spans, (spans,))
            repaired_table = Document(path).tables[0]
            self.assertEqual(len(repaired_table._tbl.tblGrid), 3)
            self.assertEqual(
                self._row_grid_spans(repaired_table),
                ((1, 2), (2, 1)),
            )

    def test_safe_horizontal_merge_moves_complete_styled_paragraph(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "styled-donor.docx"
            document, table = self._create_table_document(
                path,
                (("", "", ""),),
            )
            paragraph_style = document.styles.add_style(
                "DonorParagraph",
                WD_STYLE_TYPE.PARAGRAPH,
            )
            character_style = document.styles.add_style(
                "DonorCharacter",
                WD_STYLE_TYPE.CHARACTER,
            )
            donor = table.cell(0, 1)
            donor.text = ""
            paragraph = donor.paragraphs[0]
            paragraph.style = paragraph_style
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(7)
            first_run = paragraph.add_run("Styled")
            first_run.style = character_style
            first_run.bold = True
            first_run.font.size = Pt(13)
            first_run.font.color.rgb = RGBColor(18, 52, 86)
            second_run = paragraph.add_run(" donor")
            second_run.italic = True
            document.save(path)
            matrix = (("Styled donor", None, None),)
            spans = ((0, 0, 1, 3),)

            changed = self._repair_one(
                path,
                shape=(1, 3),
                matrix=matrix,
                spans=spans,
                widths=(100.0, 100.0, 100.0),
            )

            self.assertTrue(changed)
            repaired = Document(path)
            survivor = repaired.tables[0].cell(0, 0)
            content_paragraphs = [
                item for item in survivor.paragraphs if item.text
            ]
            self.assertEqual(len(content_paragraphs), 1)
            repaired_paragraph = content_paragraphs[0]
            self.assertEqual(repaired_paragraph.text, "Styled donor")
            self.assertEqual(repaired_paragraph.style.name, "DonorParagraph")
            self.assertEqual(
                repaired_paragraph.alignment,
                WD_ALIGN_PARAGRAPH.CENTER,
            )
            self.assertAlmostEqual(
                repaired_paragraph.paragraph_format.space_after.pt,
                7.0,
            )
            self.assertEqual(
                tuple(run.text for run in repaired_paragraph.runs),
                ("Styled", " donor"),
            )
            styled_run, italic_run = repaired_paragraph.runs
            self.assertEqual(styled_run.style.name, "DonorCharacter")
            self.assertTrue(styled_run.bold)
            self.assertAlmostEqual(styled_run.font.size.pt, 13.0)
            self.assertEqual(
                styled_run.font.color.rgb,
                RGBColor(18, 52, 86),
            )
            self.assertTrue(italic_run.italic)
            summary = inspect_editable_docx_tables(path)
            self.assertEqual(summary.table_cell_matrices, (matrix,))
            self.assertEqual(summary.table_cell_spans, (spans,))

    def test_multiple_nonempty_donors_are_rejected_atomically(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "multiple-donors.docx"
            self._create_table_document(path, (("", "Alpha", "Beta"),))

            self._assert_unchanged_after_error(
                self,
                path,
                lambda: self._repair_one(
                    path,
                    shape=(1, 3),
                    matrix=(("AlphaBeta", None, None),),
                    spans=((0, 0, 1, 3),),
                    widths=(100.0, 100.0, 100.0),
                ),
            )

    def test_text_mismatch_is_rejected_atomically(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "text-mismatch.docx"
            self._create_table_document(path, (("", "Actual"),))

            self._assert_unchanged_after_error(
                self,
                path,
                lambda: self._repair_one(
                    path,
                    shape=(1, 2),
                    matrix=(("Expected", None),),
                    spans=((0, 0, 1, 2),),
                    widths=(100.0, 100.0),
                ),
            )

    def test_drawing_field_and_nested_table_are_rejected_atomically(self):
        def add_drawing(cell):
            cell.paragraphs[0].runs[0]._r.append(OxmlElement("w:drawing"))

        def add_field(cell):
            field = OxmlElement("w:fldChar")
            field.set(qn("w:fldCharType"), "begin")
            cell.paragraphs[0].runs[0]._r.append(field)

        def add_nested_table(cell):
            cell.add_table(rows=1, cols=1)

        for name, add_hazard in (
            ("drawing", add_drawing),
            ("field", add_field),
            ("nested-table", add_nested_table),
        ):
            with self.subTest(hazard=name), tempfile.TemporaryDirectory() as temp_dir:
                path = Path(temp_dir) / f"{name}.docx"
                document, table = self._create_table_document(
                    path,
                    (("", "Value"),),
                )
                add_hazard(table.cell(0, 1))
                document.save(path)

                self._assert_unchanged_after_error(
                    self,
                    path,
                    lambda: self._repair_one(
                        path,
                        shape=(1, 2),
                        matrix=(("Value", None),),
                        spans=((0, 0, 1, 2),),
                        widths=(100.0, 100.0),
                    ),
                )

    def test_conflicting_cell_styles_are_rejected_atomically(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "style-conflict.docx"
            document, table = self._create_table_document(path, (("", ""),))
            for cell, fill in zip(table.rows[0].cells, ("FF0000", "0000FF")):
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), fill)
                cell._tc.get_or_add_tcPr().append(shading)
            document.save(path)

            self._assert_unchanged_after_error(
                self,
                path,
                lambda: self._repair_one(
                    path,
                    shape=(1, 2),
                    matrix=(("", None),),
                    spans=((0, 0, 1, 2),),
                    widths=(100.0, 100.0),
                ),
            )

    def test_equidistant_boundary_mapping_is_rejected_atomically(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "ambiguous-boundary.docx"
            self._create_table_document(
                path,
                (("A", "B"),),
                grid_widths=(1500, 1500),
                row_widths=((1500, 1500),),
            )

            self._assert_unchanged_after_error(
                self,
                path,
                lambda: self._repair_one(
                    path,
                    shape=(1, 3),
                    matrix=(("A", "B", None),),
                    spans=((0, 0, 1, 1), (0, 1, 1, 2)),
                    widths=(100.0, 100.0, 100.0),
                ),
            )

    def test_vertical_merge_restart_and_continuations_are_rebuilt_exactly(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "vertical-merge.docx"
            self._create_table_document(
                path,
                (("Group", "A"), ("", "B"), ("", "C")),
            )
            matrix = (
                ("Group", "A"),
                (None, "B"),
                (None, "C"),
            )
            spans = (
                (0, 0, 3, 1),
                (0, 1, 1, 1),
                (1, 1, 1, 1),
                (2, 1, 1, 1),
            )

            changed = self._repair_one(
                path,
                shape=(3, 2),
                matrix=matrix,
                spans=spans,
                widths=(100.0, 100.0),
            )

            self.assertTrue(changed)
            repaired_table = Document(path).tables[0]
            merge_values = []
            for row in repaired_table._tbl.findall(qn("w:tr")):
                first_cell = row.findall(qn("w:tc"))[0]
                vertical_merge = first_cell.find(
                    "./w:tcPr/w:vMerge",
                    first_cell.nsmap,
                )
                self.assertIsNotNone(vertical_merge)
                merge_values.append(vertical_merge.get(qn("w:val"), "continue"))
            self.assertEqual(tuple(merge_values), ("restart", "continue", "continue"))
            summary = inspect_editable_docx_tables(path)
            self.assertEqual(summary.table_shapes, ((3, 2),))
            self.assertEqual(summary.table_cell_counts, (4,))
            self.assertEqual(summary.table_cell_matrices, (matrix,))
            self.assertEqual(summary.table_cell_spans, (spans,))

    def test_expands_uniquely_proven_vertical_merge_paragraphs(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "aggregated-vertical-merge.docx"
            self._create_aggregated_vertical_merge(path, split_runs=True)
            matrix = (("R1", "1.5", "81"), ("R2", "3.0", "86"))
            spans = self._unit_spans(2, 3)

            changed = self._repair_one(
                path,
                shape=(2, 3),
                matrix=matrix,
                spans=spans,
                widths=(100.0, 100.0, 100.0),
                border_edges=((True, True, True, True),) * len(spans),
            )

            self.assertTrue(changed)
            summary = inspect_editable_docx_tables(path)
            self.assertEqual(summary.table_shapes, ((2, 3),))
            self.assertEqual(summary.table_cell_counts, (6,))
            self.assertEqual(summary.table_cell_spans, (spans,))
            self.assertEqual(summary.table_cell_matrices, (matrix,))
            repaired = Document(path).tables[0]
            for row in range(2):
                for column in range(3):
                    self.assertEqual(
                        self._raw_cell_border_values(repaired, row, column),
                        ("single", "single", "single", "single"),
                    )
            self.assertTrue(repaired.cell(0, 1).paragraphs[0].runs[0].bold)
            self.assertTrue(repaired.cell(0, 2).paragraphs[0].runs[0].italic)
            self.assertTrue(repaired.cell(1, 1).paragraphs[0].runs[0].underline)
            self.assertAlmostEqual(
                repaired.cell(1, 2).paragraphs[0].runs[0].font.size.pt,
                12.0,
            )

    def test_word_pagination_marker_does_not_block_vertical_expansion(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "pagination-marker.docx"
            self._create_aggregated_vertical_merge(path, split_runs=True)
            document = Document(path)
            merged_cell = document.tables[0].cell(0, 1)
            marker = OxmlElement("w:lastRenderedPageBreak")
            merged_cell.paragraphs[1].runs[0]._r.insert(0, marker)
            document.save(path)

            matrix = (("R1", "1.5", "81"), ("R2", "3.0", "86"))
            spans = self._unit_spans(2, 3)
            changed = self._repair_one(
                path,
                shape=(2, 3),
                matrix=matrix,
                spans=spans,
                widths=(100.0, 100.0, 100.0),
            )

            self.assertTrue(changed)
            summary = inspect_editable_docx_tables(path)
            self.assertEqual(summary.table_shapes, ((2, 3),))
            self.assertEqual(summary.table_cell_counts, (6,))
            self.assertEqual(summary.table_cell_spans, (spans,))
            self.assertEqual(summary.table_cell_matrices, (matrix,))
            with ZipFile(path, "r") as archive:
                document_xml = archive.read("word/document.xml")
            self.assertNotIn(b"lastRenderedPageBreak", document_xml)

    def test_systematic_run_borders_require_complete_bordered_cells(self):
        cases = (
            ("systematic", 4, None, 0, False),
            ("below-threshold", 3, None, 3, False),
            ("incomplete-cell-border", 4, 2, 4, False),
            ("nil-cell-border", 4, None, 4, True),
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            for (
                name,
                border_count,
                incomplete_row,
                expected_count,
                nil_horizontal,
            ) in cases:
                with self.subTest(name=name):
                    path = Path(temp_dir) / f"{name}.docx"
                    matrix, spans, widths = self._create_cleanup_candidate(
                        path,
                        border_count=border_count,
                        incomplete_border_row=incomplete_row,
                        nil_horizontal_borders=nil_horizontal,
                    )
                    changed = self._repair_one(
                        path,
                        shape=(4, 1),
                        matrix=matrix,
                        spans=spans,
                        widths=widths,
                    )

                    self.assertTrue(changed)
                    with ZipFile(path, "r") as archive:
                        root = etree.fromstring(
                            archive.read("word/document.xml")
                        )
                    run_borders = root.findall(
                        f".//{qn('w:rPr')}/{qn('w:bdr')}"
                    )
                    self.assertEqual(len(run_borders), expected_count)

    def test_complete_positive_row_heights_and_full_page_table_are_tightened(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "full-page-table.docx"
            matrix, spans, widths = self._create_cleanup_candidate(
                path,
                row_height=2500,
                total_width=10000,
            )

            changed = self._repair_one(
                path,
                shape=(4, 1),
                matrix=matrix,
                spans=spans,
                widths=widths,
            )

            self.assertTrue(changed)
            with ZipFile(path, "r") as archive:
                root = etree.fromstring(archive.read("word/document.xml"))
            table = root.find(f".//{qn('w:tbl')}")
            heights = table.findall(
                f"./{qn('w:tr')}/{qn('w:trPr')}/{qn('w:trHeight')}"
            )
            self.assertEqual(len(heights), 4)
            self.assertTrue(
                all(height.get(qn("w:hRule")) == "exact" for height in heights)
            )
            cell_margins = table.find(
                f"./{qn('w:tblPr')}/{qn('w:tblCellMar')}"
            )
            self.assertIsNotNone(cell_margins)
            for edge_name in ("top", "bottom"):
                edge = cell_margins.find(qn(f"w:{edge_name}"))
                self.assertIsNotNone(edge)
                self.assertEqual(edge.get(qn("w:w")), "0")
                self.assertEqual(edge.get(qn("w:type")), "dxa")
            indent = table.find(f"./{qn('w:tblPr')}/{qn('w:tblInd')}")
            self.assertIsNotNone(indent)
            self.assertEqual(indent.get(qn("w:w")), "0")
            self.assertEqual(indent.get(qn("w:type")), "dxa")
            page_margins = root.find(
                f"./{qn('w:body')}/{qn('w:sectPr')}/{qn('w:pgMar')}"
            )
            self.assertEqual(page_margins.get(qn("w:top")), "160")
            self.assertEqual(page_margins.get(qn("w:bottom")), "160")
            self.assertEqual(page_margins.get(qn("w:left")), "1120")
            self.assertEqual(page_margins.get(qn("w:right")), "1120")

    def test_page_tightening_requires_both_width_and_height_thresholds(self):
        cases = (
            ("narrow", 9000, 2500),
            ("short", 10000, 2000),
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            for name, total_width, row_height in cases:
                with self.subTest(name=name):
                    path = Path(temp_dir) / f"{name}.docx"
                    matrix, spans, widths = self._create_cleanup_candidate(
                        path,
                        row_height=row_height,
                        total_width=total_width,
                    )
                    changed = self._repair_one(
                        path,
                        shape=(4, 1),
                        matrix=matrix,
                        spans=spans,
                        widths=widths,
                    )

                    self.assertTrue(changed)
                    with ZipFile(path, "r") as archive:
                        root = etree.fromstring(
                            archive.read("word/document.xml")
                        )
                    table = root.find(f".//{qn('w:tbl')}")
                    self.assertIsNone(
                        table.find(f"./{qn('w:tblPr')}/{qn('w:tblInd')}")
                    )
                    page_margins = root.find(
                        f"./{qn('w:body')}/{qn('w:sectPr')}/{qn('w:pgMar')}"
                    )
                    self.assertEqual(page_margins.get(qn("w:top")), "1440")
                    self.assertEqual(page_margins.get(qn("w:bottom")), "1440")
                    self.assertEqual(page_margins.get(qn("w:left")), "1800")
                    self.assertEqual(page_margins.get(qn("w:right")), "1800")

    def test_row_heights_remain_flexible_when_any_row_height_is_missing(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "partial-row-heights.docx"
            matrix, spans, widths = self._create_cleanup_candidate(
                path,
                row_height=2000,
            )
            document = Document(path)
            properties = document.tables[0].rows[-1]._tr.get_or_add_trPr()
            height = properties.find(qn("w:trHeight"))
            properties.remove(height)
            document.save(path)

            changed = self._repair_one(
                path,
                shape=(4, 1),
                matrix=matrix,
                spans=spans,
                widths=widths,
            )

            self.assertTrue(changed)
            with ZipFile(path, "r") as archive:
                root = etree.fromstring(archive.read("word/document.xml"))
            heights = root.findall(f".//{qn('w:trHeight')}")
            self.assertEqual(len(heights), 3)
            self.assertTrue(
                all(height.get(qn("w:hRule")) != "exact" for height in heights)
            )

    def test_aggregated_vertical_merge_without_run_boundaries_is_rejected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "unsplittable-vertical-merge.docx"
            self._create_aggregated_vertical_merge(path, split_runs=False)

            self._assert_unchanged_after_error(
                self,
                path,
                lambda: self._repair_one(
                    path,
                    shape=(2, 3),
                    matrix=(("R1", "1.5", "81"), ("R2", "3.0", "86")),
                    spans=self._unit_spans(2, 3),
                    widths=(100.0, 100.0, 100.0),
                ),
            )
    def test_repaired_mixed_topology_matches_inspector_exactly(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "mixed-topology.docx"
            self._create_table_document(
                path,
                (
                    ("Header", "", "Score"),
                    ("Group", "Alice", "9"),
                    ("", "Bob", "8"),
                ),
            )
            matrix = (
                ("Header", None, "Score"),
                ("Group", "Alice", "9"),
                (None, "Bob", "8"),
            )
            spans = (
                (0, 0, 1, 2),
                (0, 2, 1, 1),
                (1, 0, 2, 1),
                (1, 1, 1, 1),
                (1, 2, 1, 1),
                (2, 1, 1, 1),
                (2, 2, 1, 1),
            )

            changed = self._repair_one(
                path,
                shape=(3, 3),
                matrix=matrix,
                spans=spans,
                widths=(100.0, 100.0, 100.0),
            )

            self.assertTrue(changed)
            summary = inspect_editable_docx_tables(path)
            self.assertEqual(summary.table_shapes, ((3, 3),))
            self.assertEqual(summary.table_cell_counts, (7,))
            self.assertEqual(summary.cell_count, 8)
            self.assertEqual(summary.table_cell_spans, (spans,))
            self.assertEqual(summary.table_cell_matrices, (matrix,))

    def test_repair_preserves_ignorable_namespace_declarations(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "ignorable-prefixes.docx"
            self._create_table_document(path, (("", "Value"),))
            patched_path = path.with_name("ignorable-prefixes-patched.docx")
            prefixes = {
                "cxalpha": "urn:pdf-word-converter:test:alpha",
                "cxbeta": "urn:pdf-word-converter:test:beta",
                "cxgamma": "urn:pdf-word-converter:test:gamma",
            }
            ignorable = "w14 wp14 " + " ".join(prefixes)
            declarations = "".join(
                f' xmlns:{prefix}="{namespace}"'
                for prefix, namespace in prefixes.items()
            )
            with ZipFile(path, "r") as source, ZipFile(
                patched_path,
                "w",
            ) as target:
                for info in source.infolist():
                    data = source.read(info.filename)
                    if info.filename == "word/document.xml":
                        xml = data.decode("utf-8")
                        self.assertIn('mc:Ignorable="w14 wp14"', xml)
                        xml = xml.replace(
                            "<w:document ",
                            f"<w:document{declarations} ",
                            1,
                        ).replace(
                            'mc:Ignorable="w14 wp14"',
                            f'mc:Ignorable="{ignorable}"',
                            1,
                        )
                        for prefix in prefixes:
                            self.assertEqual(xml.count(prefix), 2)
                            self.assertNotIn(f"<{prefix}:", xml)
                            self.assertNotIn(f" {prefix}:", xml)
                        data = xml.encode("utf-8")
                    target.writestr(info, data)
            patched_path.replace(path)

            changed = self._repair_one(
                path,
                shape=(1, 2),
                matrix=(("Value", None),),
                spans=((0, 0, 1, 2),),
                widths=(100.0, 100.0),
            )

            self.assertTrue(changed)
            with ZipFile(path, "r") as archive:
                repaired_xml = archive.read("word/document.xml")
            root = etree.fromstring(repaired_xml)
            mc_namespace = (
                "http://schemas.openxmlformats.org/markup-compatibility/2006"
            )
            self.assertEqual(
                root.get(f"{{{mc_namespace}}}Ignorable"),
                ignorable,
            )
            repaired_text = repaired_xml.decode("utf-8")
            for prefix, namespace in prefixes.items():
                self.assertIn(f"xmlns:{prefix}=", repaired_text)
                self.assertEqual(root.nsmap.get(prefix), namespace)
            reopened = Document(path)
            self.assertEqual(reopened.tables[0].cell(0, 0).text, "Value")
            self.assertEqual(root.tag, qn("w:document"))


if __name__ == "__main__":
    unittest.main()
